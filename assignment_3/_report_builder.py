"""Build / extend Report_Analisi_Replica.docx by APPENDING sections.

Usage:
    py _report_builder.py <section_id>

Each section is a self-contained function `section_<id>(doc)` that adds
content to the existing document. The script never rewrites earlier
sections. If the document does not yet exist, it is initialised with the
title and a short preface, then the requested section is appended.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
DOCX_PATH = ROOT / "Report_Analisi_Replica.docx"


# ---------- styling helpers --------------------------------------------------

def _set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            if len(it) == 2:
                lead, rest = it
                run = p.add_run(lead)
                run.bold = True
                p.add_run(rest)
            elif len(it) == 1:
                p.add_run(it[0])
            else:
                p.add_run(" ".join(map(str, it)))
        else:
            p.add_run(it)


def add_flag_box(doc, title, items):
    """Render a 'Flag per il prosieguo' shaded block."""
    add_para(doc, title, bold=True)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, "FFF4CE")
    cell.text = ""
    for it in items:
        p = cell.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            lead, rest = it
            r = p.add_run(lead); r.bold = True
            p.add_run(rest)
        else:
            p.add_run(it)
    # Remove the first empty paragraph that .text=""  leaves
    if cell.paragraphs and cell.paragraphs[0].text == "":
        p0 = cell.paragraphs[0]
        p0._element.getparent().remove(p0._element)


def open_methodology_box(doc, title):
    """Open a methodological deep-dive shaded cell and return the cell.
    Caller populates it with cell.add_paragraph() / add_run() as needed."""
    add_para(doc, title, bold=True)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, "E7F0FB")
    # remove the implicit empty paragraph
    if cell.paragraphs and cell.paragraphs[0].text == "":
        p0 = cell.paragraphs[0]
        p0._element.getparent().remove(p0._element)
    return cell


def box_para(cell, text, bold=False, italic=False, size=10):
    p = cell.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def box_subhead(cell, text):
    p = cell.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return p


def box_bullets(cell, items):
    for it in items:
        p = cell.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            if len(it) == 2:
                lead, rest = it
                r = p.add_run(lead); r.bold = True
                p.add_run(rest)
            elif len(it) == 1:
                p.add_run(it[0])
            else:
                p.add_run(" ".join(map(str, it)))
        else:
            p.add_run(it)


def add_table(doc, header, rows, col_widths_cm=None):
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    return p


def add_figure(doc, path, caption=None, width_cm=15.5):
    """Insert an image with optional centered caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return p


# ---------- document scaffolding --------------------------------------------

def ensure_document() -> Document:
    if DOCX_PATH.exists():
        return Document(DOCX_PATH)
    doc = Document()
    # Title page
    h = doc.add_heading("Investment Replica — Report di analisi", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    add_para(
        doc,
        "Documento di analisi sezione per sezione del notebook "
        "Portfolio_ReplicaPoliMI_v3.ipynb. Ogni sezione descrive cosa è stato "
        "fatto, con quale metodo, per quale motivo, e quali conclusioni sono "
        "state tratte. Le flag esplicite alla fine di ogni sezione raccolgono "
        "le considerazioni da riutilizzare nelle sezioni successive.",
        italic=True,
    )
    doc.add_page_break()
    return doc


# ---------- SECTION 1: Esplorazione e costruzione del target -----------------

def section_1(doc):
    add_heading(doc, "1. Esplorazione dei dati e costruzione del target", level=1)

    # --- 1.1 Dataset ---
    add_heading(doc, "1.1 Il dataset", level=2)
    add_para(
        doc,
        "Il dataset Bloomberg fornito (Dataset3_PortfolioReplicaStrategy.xlsx) "
        "contiene 705 osservazioni settimanali in valuta locale che coprono "
        "il periodo 23 ottobre 2007 – 20 aprile 2021. Le serie sono di due "
        "categorie:",
    )
    add_bullets(
        doc,
        [
            ("Quattro indici target candidati: ",
             "MXWO (MSCI World), MXWD (MSCI ACWI), LEGATRUU (Bloomberg Global "
             "Aggregate Bond) e HFRXGL (HFRX Global Hedge Fund)."),
            ("Undici futures replicanti: ",
             "ES1, NQ1, VG1, TP1, LLL1 (equity dei vari mercati), "
             "RX1, TY1, DU1, TU2 (bond US e tedeschi a varie scadenze), "
             "GC1 (oro), CO1 (Brent)."),
        ],
    )
    add_para(
        doc,
        "Il periodo include eventi di stress significativi: la crisi "
        "finanziaria 2008–2009, il flash crash 2010, la crisi del debito "
        "sovrano EU 2011, il sell-off cinese 2015, il Q4 2018 e il crollo "
        "Covid del 2020. Sono esattamente le finestre in cui un replicatore "
        "deve dimostrare di tenere; tenerle nel sample è cruciale.",
    )

    # --- 1.2 Diagnostica delle serie ---
    add_heading(doc, "1.2 Statistiche e diagnostica delle serie", level=2)
    add_para(
        doc,
        "Sui returns settimanali dei quattro indici target abbiamo calcolato "
        "rendimento e volatilità annualizzati, Sharpe, max drawdown, skewness "
        "e kurtosis. I numeri rilevanti:",
    )
    add_table(
        doc,
        header=["Serie", "Ann. Ret", "Ann. Vol", "Sharpe", "Max DD", "Skew", "Kurt"],
        rows=[
            ["MXWO",     "5.90%",  "17.52%", "0.34",  "−57.6%", "−0.68", "4.36"],
            ["MXWD",     "5.51%",  "17.78%", "0.31",  "−58.4%", "−0.64", "4.66"],
            ["LEGATRUU", "3.37%",  "5.34%",  "0.63",  "−10.4%", "−0.33", "4.12"],
            ["HFRXGL",   "0.50%",  "4.78%",  "0.10",  "−24.9%", "−2.30", "13.70"],
        ],
    )
    add_para(
        doc,
        "Tre osservazioni chiave: (i) HFRX ha kurtosis 13.7 e skewness −2.3 — "
        "una distribuzione fortemente asimmetrica e a code spesse che esclude "
        "l'utilizzo di un VaR Gaussiano; (ii) tutti gli indici hanno skewness "
        "negativa, ovvero le perdite sono più estreme dei guadagni; (iii) il "
        "bond aggregato (LEGATRUU) è la serie più \"benigna\" in termini di "
        "Sharpe e drawdown.",
    )

    add_para(
        doc,
        "Per la replica conta anche il downside: a Sharpe abbiamo aggiunto "
        "Sortino (excess return su downside deviation) e Calmar (return su "
        "max drawdown). HFRX risulta debole su entrambi (Sortino 0.10, "
        "Calmar 0.02), confermando che la sua bassa volatilità è un artefatto "
        "del fatto che gli hedge fund index riportano rendimenti smussati.",
    )

    # --- 1.3 Stazionarietà ---
    add_heading(doc, "1.3 Stazionarietà: ADF e KPSS", level=2)
    add_para(
        doc,
        "Una regressione lineare sui returns presuppone (debole) "
        "stazionarietà. Per validare questa ipotesi e — soprattutto — per "
        "escludere l'errore comune di regredire sui prezzi, abbiamo eseguito "
        "Augmented Dickey-Fuller (H₀: presenza di radice unitaria) e KPSS "
        "(H₀: serie stazionaria) su ciascun target sia in livelli che in "
        "returns. Esito sintetico:",
    )
    add_bullets(
        doc,
        [
            ("Prezzi: ", "ADF p ≈ 0.99 per MXWO e MXWD, 0.67 per LEGATRUU, "
             "0.17 per HFRX → non si rigetta la radice unitaria. KPSS rigetta "
             "la stazionarietà su tutte e quattro. Coerente: i prezzi non "
             "sono stazionari."),
            ("Returns: ", "ADF p < 10⁻⁴ su tutte le serie (output saturato a "
             "0.0000), KPSS p ≥ 0.10 → le serie di returns sono stazionarie. "
             "Il warning di statsmodels (\"actual p-value is smaller than "
             "the p-value returned\") conferma che la statistica test cade "
             "fuori dalla tabella di lookup: per n = 704 osservazioni quasi "
             "i.i.d., il test ha potenza praticamente unitaria."),
        ],
    )
    add_para(
        doc,
        "Conclusione operativa: tutte le successive regressioni e backtest "
        "saranno condotti sui returns, mai sui prezzi.",
    )

    # --- 1.4 Costruzione del target ---
    add_heading(doc, "1.4 Costruzione del target — Monster Index", level=2)
    add_para(
        doc,
        "Il target da replicare è una combinazione lineare a pesi fissi degli "
        "indici componente:",
    )
    add_formula(doc, "Monster = 0.50 · HFRXGL + 0.25 · MXWO + 0.25 · LEGATRUU")
    add_para(
        doc,
        "La scelta è coerente con un'asset allocation \"endowment-style\": "
        "metà alternative, un quarto equity, un quarto bond. L'esposizione "
        "dominante è hedge fund, che è anche il pezzo più difficile da "
        "replicare con futures (gli hedge fund usano leva, derivati esotici, "
        "skill manageriale). I returns della componente HFRX vengono usati "
        "in forma raw — un eventuale unsmoothing era stato testato ma "
        "ridistribuiva varianza in modo che peggiorava la replica empirica.",
    )
    add_para(
        doc,
        "Statistiche del target costruito (704 osservazioni, 30 ott 2007 → "
        "20 apr 2021):",
    )
    add_table(
        doc,
        header=["Metrica", "Valore"],
        rows=[
            ["Ann. Return",   "2.57%"],
            ["Ann. Volatility", "6.33%"],
            ["Sharpe",        "0.41"],
            ["Max Drawdown",  "−29.0%"],
            ["Skewness",      "−1.35"],
            ["Kurtosis",      "7.64"],
        ],
    )
    add_para(
        doc,
        "Il target eredita la natura fat-tailed e left-skewed dei suoi "
        "componenti — kurtosis 7.6 è ben oltre il valore Gaussiano (3). "
        "Questo è un input critico per la scelta della misura di rischio "
        "nel backtest engine.",
    )

    # --- 1.5 Target vs futures ---
    add_heading(doc, "1.5 Relazione target ↔ futures", level=2)
    add_para(
        doc,
        "La correlazione del target con ciascun future settimanale dà una "
        "prima indicazione su quali strumenti porteranno informazione utile "
        "alla regressione:",
    )
    add_table(
        doc,
        header=["Future", "Asset class", "Corr con target"],
        rows=[
            ["ES1",  "Equity DM (S&P 500)",      "+0.838"],
            ["NQ1",  "Equity DM (Nasdaq 100)",   "+0.747"],
            ["VG1",  "Equity DM (Eurostoxx 50)", "+0.731"],
            ["TP1",  "Equity DM (Topix)",        "+0.599"],
            ["LLL1", "Equity EM",                "+0.499"],
            ["CO1",  "Commodity (Brent)",        "+0.445"],
            ["DU1",  "Bond (Schatz 2y)",         "−0.262"],
            ["GC1",  "Commodity (Oro)",          "+0.218"],
            ["TU2",  "Bond (UST 2y)",            "−0.153"],
            ["RX1",  "Bond (Bund 10y)",          "−0.131"],
            ["TY1",  "Bond (UST 10y)",           "−0.119"],
        ],
    )
    add_para(
        doc,
        "La replica sarà dominata dall'equity beta (correlazioni > 0.7 sui "
        "primi tre); le commodities hanno un contributo intermedio (Brent "
        "0.45 grazie alla correlazione con il ciclo macro); i bond sono "
        "debolmente, e in maggioranza negativamente, correlati al target — "
        "il loro ruolo nella replica sarà di hedge / diversificazione "
        "anziché di driver di rendimento.",
    )

    # --- 1.6 Correlazioni condizionali ---
    add_heading(doc, "1.6 Correlazioni condizionali ai regimi di volatilità", level=2)
    add_para(
        doc,
        "Le matrici di correlazione vengono calcolate tre volte: sul "
        "campione intero, sui sottoperiodi con volatilità rolling 26 "
        "settimane dell'equity (MXWO) sotto la mediana (\"calm\") e sopra "
        "l'80° percentile (\"stress\"). Il fenomeno noto è che in regime di "
        "stress le correlazioni cross-asset si compattano: la diversificazione "
        "viene meno proprio quando servirebbe.",
    )
    add_para(
        doc,
        "Lettura qualitativa: nel regime calm le correlazioni positive "
        "interne all'equity restano dominanti, e i bond mantengono "
        "correlazioni vicine a zero o leggermente negative col target. Nel "
        "regime stress le correlazioni equity-target si comprimono ancora "
        "verso 1, le correlazioni inter-bond restano stabili, ma la "
        "correlazione bond-equity tende ad aumentare in valore assoluto. "
        "Significato: un replicatore con pesi statici (calibrati su tutto il "
        "sample) sotto-pesa la dinamica dei regimi.",
    )

    # --- 1.7 Autocorrelazione e clustering ---
    add_heading(doc, "1.7 Autocorrelazione e volatility clustering", level=2)
    add_para(
        doc,
        "Test di Ljung-Box sui returns a 5/10/20 lag e ispezione di "
        "ACF/PACF. Risultato sintetico:",
    )
    add_table(
        doc,
        header=["Serie", "Ljung-Box p (lag 5)", "Ljung-Box p (lag 20)"],
        rows=[
            ["MXWO",     "0.591", "0.052"],
            ["MXWD",     "0.539", "0.043"],
            ["LEGATRUU", "0.136", "0.148"],
            ["HFRXGL",   "0.000", "0.000"],
        ],
    )
    add_para(
        doc,
        "I tre indici tradizionali (equity e bond) sono indistinguibili "
        "dal white noise per quanto riguarda la media condizionale — "
        "compatibile con una replica via regressione lineare statica nella "
        "media (cioè senza componente AR esplicita). HFRX è l'unica serie "
        "fortemente autocorrelata, fenomeno noto come return smoothing "
        "tipico degli indici hedge fund (illiquidità, mark-to-model dei "
        "manager, return reporting con ritardo). Avendo deciso di costruire "
        "il target con HFRX raw e non con un unsmoothing, ci aspettiamo che "
        "qualunque replicatore lineare lasci un residuo strutturale legato "
        "a quella persistenza.",
    )
    add_para(
        doc,
        "L'ACF dei returns al quadrato è invece significativa per tutte le "
        "serie: c'è chiaro volatility clustering. Questo è il pezzo che "
        "motiva un VaR non-Gaussiano (Cornish-Fisher) e — più avanti — la "
        "scaling adattiva del portafoglio in base alla volatilità storica.",
    )

    # --- Approfondimento metodologico: ACF / PACF / Ljung-Box ---
    box = open_methodology_box(
        doc, "Approfondimento metodologico — ACF, PACF, Ljung-Box"
    )

    box_subhead(box, "Autocorrelazione")
    box_para(
        box,
        "Una serie storica r_t è autocorrelata se i suoi valori a tempi "
        "diversi non sono indipendenti: il passato porta informazione sul "
        "futuro. L'autocorrelazione al lag k è la correlazione fra la "
        "serie e una sua copia traslata di k periodi:",
    )
    box_para(box, "    ρ(k) = Corr(r_t, r_{t−k})", italic=True)
    box_para(
        box,
        "ρ(0) vale sempre 1; ρ(k) ∈ [−1, +1] per k > 0; se ρ(k) ≈ 0 per "
        "ogni k > 0 la serie è indistinguibile dal white noise nella media.",
    )

    box_subhead(box, "ACF — Autocorrelation Function")
    box_para(
        box,
        "L'ACF è il grafico di ρ(k) al variare di k. Sull'asse x compare il "
        "lag in settimane, sull'asse y il valore di ρ(k) rappresentato da "
        "barre verticali. Le bande blu ombreggiate sono l'intervallo di "
        "confidenza al 95% sotto l'ipotesi nulla \"serie = white noise\"; "
        "per n ≈ 700 osservazioni la banda è circa ±1.96/√n ≈ ±0.074. "
        "Barre dentro la banda → non significativamente diverse da zero; "
        "barre che bucano la banda → autocorrelazione significativa a "
        "quel lag.",
    )
    box_bullets(
        box,
        [
            ("Lettura nei nostri grafici (MXWO, MXWD, LEGATRUU): ",
             "tutte le barre dei primi 20 lag restano dentro la banda → "
             "compatibili con white noise → la regressione lineare statica "
             "nella media è giustificata."),
            ("Lettura per HFRX: ",
             "i primi 2–3 lag dell'ACF bucano chiaramente la banda, con "
             "decadimento lento. Questo è il return smoothing tipico degli "
             "hedge fund index."),
        ],
    )

    box_subhead(box, "PACF — Partial Autocorrelation Function")
    box_para(
        box,
        "L'ACF mostra la correlazione \"totale\" fra r_t e r_{t−k}, "
        "inclusa la parte indiretta che passa attraverso i lag intermedi. "
        "La PACF isola la sola componente diretta: la correlazione fra r_t "
        "e r_{t−k} al netto dell'effetto di r_{t−1}, …, r_{t−k+1}. "
        "Formalmente è il coefficiente φ_kk nella regressione AR(k):",
    )
    box_para(box, "    r_t = φ_k1·r_{t−1} + … + φ_kk·r_{t−k} + ε_t",
             italic=True)
    box_para(
        box,
        "Uso pratico: la PACF è lo strumento standard per identificare "
        "l'ordine di un AR(p). Se la PACF si annulla bruscamente dopo il "
        "lag p (barre dentro la banda da p+1 in poi) → suggerisce un "
        "AR(p). Se invece è l'ACF che si annulla bruscamente mentre la "
        "PACF decade gradualmente → suggerisce un MA(q). Per HFRX si "
        "osserva PACF(1) dominante con PACF(2+) ai limiti, struttura "
        "compatibile con un AR(1).",
    )

    box_subhead(box, "Test di Ljung-Box")
    box_para(
        box,
        "ACF e PACF sono diagnostiche visive. Il Ljung-Box è il "
        "corrispettivo test formale, valuta congiuntamente "
        "l'autocorrelazione fino al lag m con statistica:",
    )
    box_para(
        box,
        "    Q(m) = n·(n+2) · Σ_{k=1..m} ρ̂(k)² / (n − k)",
        italic=True,
    )
    box_para(
        box,
        "Sotto H₀ (no autocorrelazione fino al lag m) Q segue una "
        "distribuzione χ² con m gradi di libertà. Output operativo: "
        "p-value alto (> 0.05) → non si rigetta H₀ → serie compatibile "
        "con white noise; p-value basso → autocorrelazione presente. "
        "Nei nostri risultati MXWO/MXWD/LEGATRUU hanno p > 0.05 ai "
        "principali lag, HFRX dà p = 0.000 a lag 5, 10 e 20 — "
        "autocorrelazione strutturale a tutti gli orizzonti.",
    )

    box_subhead(box, "Volatility clustering — ACF dei returns al quadrato")
    box_para(
        box,
        "Una serie può essere non autocorrelata nella media ma "
        "autocorrelata nella varianza: settimane volatili seguono "
        "settimane volatili. Diagnostica standard: calcolare l'ACF della "
        "serie r_t² (proxy della varianza condizionale). Nei nostri "
        "grafici tutte e quattro le serie mostrano barre fuori dalla "
        "banda sui primi 5–10 lag, persino LEGATRUU. Il vol clustering "
        "è pervasivo.",
    )
    box_para(
        box,
        "Conseguenze pratiche già richiamate nelle flag: la varianza è "
        "prevedibile dalla storia recente → il VaR scaling con lookback "
        "rolling (3 anni) ha senso; la distribuzione condizionale ha "
        "code più spesse della incondizionale → Cornish-Fisher è "
        "preferibile al VaR Gaussiano.",
    )

    box_subhead(box, "Sintesi per la lettura dei grafici del notebook")
    box_bullets(
        box,
        [
            ("Confronto visivo della banda: ",
             "per MXWO, MXWD, LEGATRUU pressoché tutte le barre dell'ACF "
             "stanno dentro l'intervallo blu → nessuna autocorrelazione "
             "rilevante."),
            ("HFRX rompe il pattern: ",
             "ACF(1) ≈ 0.3, decay lento, PACF(1) dominante, p di "
             "Ljung-Box azzerato. Smoothing strutturale che il modello "
             "lineare non può eliminare."),
            ("Plot dell'ACF di r²: ",
             "barre fuori dalla banda anche su LEGATRUU → vol clustering "
             "universale → giustifica VaR rolling e Cornish-Fisher."),
        ],
    )

    # --- Flag finali ---
    flags = [
        ("HFRX domina il target (50%) ed è autocorrelato: ",
         "qualunque replicatore lineare statico lascerà un errore "
         "strutturale su quella componente. Tenere a mente quando "
         "interpretiamo la Tracking Error: parte è \"colpa\" del target, "
         "non del modello."),
        ("Code spesse e skewness negativa (Kurt = 7.6, Skew = −1.35): ",
         "il VaR Gaussiano sotto-stima il rischio nelle code. Il backtest "
         "engine userà Cornish-Fisher come default."),
        ("Equity beta è il motore della replica (ES1, NQ1, VG1 corr > 0.7): ",
         "qualunque modello regolarizzato concentrerà i pesi su questi tre, "
         "lasciando i bond come variabili di hedge marginali. Da usare come "
         "sanity check sulle weight composition dei prossimi modelli."),
        ("Correlazioni regime-dipendenti: ",
         "motiva (i) modelli adattivi/online come Kalman e (ii) approcci "
         "regime-aware come l'HMM (Idea 2) e il conformal condizionato ai "
         "regimi (Idea D2)."),
        ("Volatility clustering pervasivo: ",
         "giustifica il VaR scaling con lookback rolling (3 anni) e tutto "
         "l'apparato di risk management nel backtest engine."),
        ("Domini di stress nel sample (2008, 2011, 2020): ",
         "ogni metrica di replica sarà valutata principalmente OOS in "
         "queste finestre. Conservarle nel sample è la decisione di "
         "data-engineering più impattante dell'intero progetto."),
    ]
    add_flag_box(doc, "Flag per il prosieguo del progetto", flags)

    add_para(doc, "")  # spacer
    return doc


# ---------- dispatcher -------------------------------------------------------

# ---------- SECTION 2: Backtest engine, baselines, Elastic Net ----------------

def section_2(doc):
    add_heading(doc, "2. Unified Backtest Engine, baselines e tuning Elastic Net",
                level=1)

    # 2.1 Backtest engine
    add_heading(doc, "2.1 Architettura del backtest engine", level=2)
    add_para(
        doc,
        "Un portafoglio è una combinazione lineare di returns di asset: "
        "r_p = Σⱼ wⱼ · rⱼ. Questo è il motivo per cui una regressione lineare "
        "(eventualmente penalizzata) è il modello statisticamente più "
        "naturale per la replica: i suoi coefficienti sono direttamente i "
        "pesi di portafoglio cercati. L'engine costruito attorno a "
        "scikit-learn rispetta quattro principi operativi:",
    )
    add_bullets(
        doc,
        [
            ("Walk-forward su finestra rolling: ",
             "ad ogni settimana t si stima sulla finestra K precedente, si "
             "applica la regola di scaling del rischio e si proietta su "
             "t+1. Nessuna informazione futura entra mai nelle stime, e le "
             "metriche sono calcolate solo OOS."),
            ("Standardizzazione (StandardScaler): ",
             "i returns settimanali hanno excess kurtosis ≈ 8; un MinMax "
             "scaler sarebbe stato dominato dall'osservazione di crisi "
             "peggiore della finestra, deformando lo scaling. Lo "
             "StandardScaler (media nulla, std unitaria) distribuisce "
             "l'effetto su tutto il campione. I coefficienti vengono "
             "riportati su scala raw via β_raw,j = β_std,j · σ_y / σ_xj."),
            ("Intercetta abilitata: ",
             "vincolando l'intercetta a zero quando il target ha media non "
             "nulla si introduce bias su tutti i coefficienti. "
             "L'intercetta non è un peso del portafoglio — la trattiamo "
             "come piccola gamba cash."),
            ("Costi di transazione e gestione del rebalance: ",
             "una flat fee di 2 bps per unità di notional turnata; la "
             "frequenza di rebalance è un parametro (settimanale di "
             "default, mensile come variante). Rebalance settimanali con "
             "modelli non regolarizzati possono erodere ~200 bps/anno."),
        ],
    )
    add_para(
        doc,
        "Default operativi adottati nel resto del notebook: rolling window "
        "156 settimane per i baseline (= 3 anni), 208 settimane (= 4 anni) "
        "per i modelli post-tuning; VaR al 99% su orizzonte 4 settimane; "
        "tetto VaR 20% (limite UCITS); VaR lookback 156 settimane.",
    )

    # 2.2 VaR
    add_heading(doc, "2.2 Misura di rischio: Cornish-Fisher VaR", level=2)
    add_para(
        doc,
        "Tre stimatori sono implementati; per chiarezza, la formula "
        "Gaussiana è:",
    )
    add_formula(doc, "VaRα(h) = − zα · σ · √h")
    add_para(
        doc,
        "dove zα è il quantile della normale standard al livello α e h è "
        "l'orizzonte in settimane. Lo stimatore Cornish-Fisher, scelto "
        "come default, è una correzione della Gaussiana che incorpora "
        "skewness sk e excess kurtosis ek dei returns:",
    )
    add_formula(
        doc,
        "z_CF = zα + (zα² − 1)·sk/6 + (zα³ − 3zα)·ek/24 "
        "− (2zα³ − 5zα)·sk²/36"
    )
    add_formula(doc, "VaR_CF(h) = − z_CF · σ · √h")
    add_para(
        doc,
        "La scelta è motivata direttamente dalle flag della Sezione 1: il "
        "Monster Index ha kurtosis 7.6 e skewness −1.35, quindi il VaR "
        "Gaussiano sottostima sistematicamente il quantile del 99%. La "
        "correzione di Cornish-Fisher è in forma chiusa, computa in O(n), "
        "e abbiamo verificato che produce un VaR del 20–30% più "
        "conservativo rispetto al Gaussiano nelle finestre con vol "
        "clustering elevato. Il rischio viene poi confrontato con il cap "
        "del 20% (limite UCITS) e i pesi vengono scalati di un fattore "
        "min(1, VAR_MAX / VaR_realizzato).",
    )

    # 2.3 Baselines
    add_heading(doc, "2.3 Baseline: OLS, Ridge, Lasso, Elastic Net (default)",
                level=2)
    add_para(
        doc,
        "Quattro modelli con default sensati e finestra 156 settimane, "
        "rebalance settimanale, allineati sullo stesso intervallo OOS. "
        "Le configurazioni: OLS senza penalty; Ridge con α = 1; Lasso "
        "con α = 10⁻³; ElasticNet con α = 10⁻³ e l1_ratio = 0.5. "
        "Risultati sintetici sull'OOS comune a tutti i modelli (496 "
        "settimane, OOS start 25 ottobre 2011):",
    )
    add_table(
        doc,
        header=["Modello", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett."],
        rows=[
            ["OLS",        "0.0279", "−0.162", "0.855", "1.72", "0.089"],
            ["Ridge",      "0.0280", "−0.160", "0.854", "1.69", "0.086"],
            ["Lasso",      "0.0279", "−0.156", "0.855", "1.67", "0.088"],
            ["ElasticNet", "0.0279", "−0.159", "0.855", "1.69", "0.088"],
        ],
    )
    add_para(
        doc,
        "I quattro modelli sono indistinguibili in termini di Tracking "
        "Error, IR e correlazione (i numeri coincidono fino alla seconda "
        "cifra decimale). La differenza vera è in due voci: la "
        "regolarizzazione blanda non riduce la gross exposure (che resta "
        "sopra 1.65, cioè il portafoglio è significativamente long-short "
        "con leva del 70%); il turnover settimanale è ≈ 9% del notional, "
        "che a 2 bps di costo produce circa 90 bps/anno di erosione. La "
        "regolarizzazione forte (alti α) e la frequenza di rebalance più "
        "bassa, oggetto del tuning Elastic Net, sono entrambe necessarie "
        "per gestire questi due numeri.",
    )

    # 2.4 EN tuning
    add_heading(doc, "2.4 Tuning di Elastic Net", level=2)
    add_para(
        doc,
        "Elastic Net minimizza la perdita quadratica regolarizzata con un "
        "mix di L1 e L2:",
    )
    add_formula(
        doc,
        "min_w  (1/2n)·‖Xw − y‖²  +  α · [ ρ·‖w‖₁ + (1 − ρ)/2 · ‖w‖²² ]"
    )
    add_para(
        doc,
        "dove α controlla l'intensità complessiva della penalty e ρ "
        "(l1_ratio) sposta il bilancio fra L1 e L2. Per ρ = 1 si recupera "
        "Lasso puro, sparsifica; per ρ = 0 si recupera Ridge puro, "
        "smussa. Valori intermedi combinano sparsità e stabilità — utili "
        "quando i regressori (i futures) sono fortemente correlati fra "
        "loro come nel nostro caso.",
    )
    add_para(
        doc,
        "Lo spazio di ricerca è 10 valori di α (1e−5 ÷ 3e−1, log-scala), "
        "6 valori di ρ (0.1, 0.3, 0.5, 0.7, 0.9, 1.0) e 4 lunghezze di "
        "finestra (104, 156, 208, 260 settimane), per 240 configurazioni "
        "totali. La validazione usa solo il primo 60% del campione, con "
        "TimeSeriesSplit (4 split), e il punteggio è la TE media OOS. Il "
        "restante 40% resta completamente held-out.",
    )
    add_para(
        doc,
        "Le prime 5 configurazioni del ranking CV TE sono di fatto "
        "indistinguibili:",
    )
    add_table(
        doc,
        header=["window", "α", "l1_ratio", "CV TE"],
        rows=[
            ["208", "0.10", "0.3", "0.002962"],
            ["208", "0.10", "0.1", "0.002975"],
            ["208", "0.10", "0.5", "0.002988"],
            ["208", "0.30", "0.1", "0.002992"],
            ["208", "0.03", "1.0", "0.002997"],
        ],
    )
    add_para(
        doc,
        "Tutto il top-5 cade su window = 208 (4 anni) e α ∈ {0.03, 0.1, "
        "0.3}, segno che la finestra di training è il parametro più "
        "rilevante e che l'intensità della regolarizzazione è "
        "sufficientemente forte da imporre stabilità ai pesi. Il "
        "vincitore è (window = 208, α = 0.10, l1_ratio = 0.3), uno "
        "schema a forte componente L2 (sparsità moderata) con una "
        "finestra lunga.",
    )
    add_para(
        doc,
        "La differenza in CV TE fra il top-1 e il top-5 è meno dell'1.2% "
        "(0.002962 vs 0.002997); la loss surface è piatta nei dintorni "
        "dell'ottimo. Questo dà solidità alla scelta — non c'è una "
        "configurazione \"miracolosa\" — ma anche flessibilità: se in "
        "futuro avessimo motivi per preferire più sparsità (es. minor "
        "numero di asset effettivamente attivi), potremmo passare a "
        "l1_ratio = 1.0 senza degradare l'errore di tracking.",
    )

    # 2.5 Schema di valutazione
    add_heading(doc, "2.5 Schema di valutazione: training, validation, test",
                level=2)
    add_para(
        doc,
        "In questa parte del notebook convivono due protocolli distinti di "
        "valutazione: uno per i baseline (iperparametri fissati a priori, "
        "nessuna validation) e uno per l'Elastic Net tunato (tuning con "
        "search/eval split + cross-validation interna). Distinguerli è "
        "essenziale per leggere correttamente i numeri riportati e per "
        "estendere correttamente il protocollo ai modelli che vedremo nelle "
        "sezioni successive.",
    )

    add_heading(doc, "Protocollo A — Baselines: rolling walk-forward senza tuning",
                level=3)
    add_para(
        doc,
        "I quattro baseline (OLS, Ridge α = 1, Lasso α = 10⁻³, EN default "
        "α = 10⁻³, ρ = 0.5) hanno iperparametri fissati a priori. Non c'è "
        "una fase di validation perché non ci sono iperparametri da "
        "scegliere. L'engine backtest() esegue uno schema walk-forward "
        "puro: per ogni settimana t a partire da t = rolling_window (= 156), "
        "stima il modello sulla finestra rolling [t−156, t), applica il "
        "VaR scaling, e proietta i pesi su t+1 generando un return OOS. "
        "Le metriche (TE, IR, Correlation) sono statistiche su questa "
        "serie di returns OOS.",
    )
    add_bullets(
        doc,
        [
            ("Training set: ", "la finestra rolling di 156 settimane "
             "antecedente t."),
            ("Test set: ", "la singola settimana t+1 (la successiva)."),
            ("Validation set: ", "non esiste — non ci sono iperparametri "
             "da selezionare."),
        ],
    )
    add_para(
        doc,
        "Per i baseline l'OOS effettivo va dalla 157ª settimana del "
        "campione fino all'ultima, ovvero 9 ottobre 2010 → 20 aprile 2021 "
        "(548 settimane).",
    )

    add_heading(doc, "Protocollo B — Elastic Net tunato: schema a tre stadi",
                level=3)
    add_para(
        doc,
        "Per il modello con iperparametri da scegliere il sample viene "
        "diviso cronologicamente una sola volta in due blocchi disgiunti:",
    )
    add_bullets(
        doc,
        [
            ("Search set (60% iniziale del sample, ≈ 422 settimane, ottobre "
             "2007 → metà 2015): ", "usato per scegliere gli iperparametri."),
            ("Eval set (40% finale, ≈ 282 settimane, metà 2015 → aprile "
             "2021): ", "completamente held-out dalla scelta degli "
             "iperparametri. Non viene mai \"visto\" dal grid search."),
        ],
    )
    add_para(
        doc,
        "All'interno del search set, ogni configurazione (window, α, "
        "l1_ratio) viene valutata tramite TimeSeriesSplit della libreria "
        "scikit-learn con n_splits = 4. TimeSeriesSplit produce quattro "
        "coppie (train_k, test_k) tali che train_k precede sempre test_k "
        "cronologicamente, e la dimensione del train cresce a ogni split. "
        "Per ciascuno split:",
    )
    add_bullets(
        doc,
        [
            "Si stima il modello sulle ultime window settimane del "
            "blocco train_k (per rispettare la rolling window scelta).",
            "Si proiettano i pesi su test_k e si calcola la Tracking "
            "Error.",
            "Si media la TE sui 4 split → score della configurazione.",
        ],
    )
    add_para(
        doc,
        "Selezionata la configurazione vincente (window = 208, α = 0.10, "
        "l1_ratio = 0.3), il backtest finale viene rieseguito sul sample "
        "completo con la stessa logica walk-forward usata per i baseline. "
        "Le metriche riportate in 2.6 e 2.7 sono dunque mediate su tutto il "
        "periodo OOS, non solo sull'eval set held-out.",
    )

    box = open_methodology_box(
        doc,
        "Nota statistica — sovrapposizione fra search set e OOS del "
        "backtest finale"
    )
    box_para(
        box,
        "Il backtest finale, eseguito sul sample completo, copre l'OOS "
        "dalla settimana 208 in poi (25 ottobre 2011 → 20 aprile 2021, "
        "496 settimane). Di queste, le prime 214 (settimane 208 → 422) "
        "cadono dentro il search set; le ultime 282 (settimane 422 → 704) "
        "sono nell'eval set truly held-out. Esiste quindi una leggera "
        "dipendenza fra i dati usati per scegliere (window, α, l1_ratio) "
        "e quelli su cui si misurano TE/IR finali.",
    )
    box_para(
        box,
        "La scelta è consapevole. Tre ragioni la motivano:",
    )
    box_bullets(
        box,
        [
            ("Impronta minima sulla scelta: ",
             "il grid search ha estratto da quei dati solamente 3 "
             "scalari (window, α, l1_ratio) tramite una statistica "
             "aggregata su 4 fold piccoli. Il rischio di overfitting "
             "agli iperparametri è basso, e la loss surface piatta "
             "(top-5 CV TE entro 1.2%) lo conferma."),
            ("Esigenza di continuità per le analisi a valle: ",
             "bootstrap CI, conformal prediction (Idea D), RL "
             "(Idea E) richiedono serie OOS lunghe e contigue. Limitarsi "
             "al 40% held-out (282 settimane) lascerebbe poco materiale "
             "per quei metodi."),
            ("Confronto su pari OOS con i baseline: ",
             "i baseline non hanno tuning e usano l'intero OOS post-156. "
             "Per metterli a confronto con l'EN tunato senza vantaggi "
             "ingiusti, è essenziale che entrambi usino lo stesso "
             "intervallo OOS — vedi protocollo C."),
        ],
    )
    box_para(
        box,
        "Sanity check effettuato: le metriche calcolate solo sulle 282 "
        "settimane truly held-out non spostano il ranking fra modelli "
        "rispetto a quelle calcolate sull'OOS pieno (496 settimane). La "
        "decisione di consegna non cambia.",
    )

    add_heading(doc, "Protocollo C — Allineamento dei confronti", level=3)
    add_para(
        doc,
        "Baselines e EN tunato hanno window diverse (156 vs 208), quindi "
        "l'OOS effettivo inizia in date diverse: 9 ottobre 2010 per i "
        "baseline, 25 ottobre 2011 per l'EN tunato. Per evitare che la "
        "metrica di un modello sia gonfiata o degradata da settimane in "
        "cui un altro modello non aveva ancora prodotto pesi, tutti i "
        "risultati riportati nella tabella di confronto vengono ri-tagliati "
        "sull'intersezione delle OOS, cioè dal 25 ottobre 2011 in poi (496 "
        "settimane). Lo stesso protocollo verrà riapplicato in tutte le "
        "sezioni successive in cui modelli con window diverse coesistono.",
    )

    # 2.6 Aligned comparison (was 2.5)
    add_heading(doc, "2.6 Confronto allineato baselines vs EN tunato",
                level=2)
    add_para(
        doc,
        "Re-runnati sull'OOS comune (496 settimane), i quattro baseline "
        "vs le due varianti tunate (weekly e monthly rebalancing):",
    )
    add_table(
        doc,
        header=["Modello", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett."],
        rows=[
            ["OLS",                              "0.0279", "−0.161", "0.855", "1.72", "0.089"],
            ["Ridge",                            "0.0280", "−0.160", "0.854", "1.69", "0.086"],
            ["Lasso",                            "0.0279", "−0.156", "0.855", "1.67", "0.088"],
            ["EN default",                       "0.0279", "−0.159", "0.855", "1.69", "0.088"],
            ["EN tuned, weekly",                 "0.0280", "−0.195", "0.854", "0.64", "≈ 0.030"],
            ["EN tuned, monthly (elnet_monthly)","0.0281", "−0.199", "0.853", "0.64", "0.018"],
        ],
    )
    add_para(
        doc,
        "Tre conclusioni operative:",
    )
    add_bullets(
        doc,
        [
            ("Cambio strutturale di profilo: ",
             "EN tunato dimezza la gross exposure (da 1.7 a 0.64). Resta "
             "ben sotto il cap UCITS del 100% e diventa "
             "operativamente compatibile con un UCITS fund. Il turnover "
             "settimanale cala di 5× passando da weekly a monthly "
             "rebalance."),
            ("Pari TE e correlazione: ",
             "TE e correlazione non si muovono di fatto fra i baseline e "
             "l'EN tunato (≈ 2.79–2.81%, corr ≈ 0.853–0.855). La "
             "regolarizzazione forte non costa accuratezza di tracking."),
            ("Il prezzo della regolarizzazione: l'IR netto peggiora ",
             "da −0.16 (baselines) a −0.20 (EN tunato monthly). La "
             "differenza è quasi tutta nella media del rendimento della "
             "replica: i baseline catturano 4.02% annualizzato, l'EN "
             "tunato monthly 3.91% — circa 11 bps/anno in meno. È il "
             "costo della sparsità: la replica diventa più povera nella "
             "coda di rendimenti che dipendono da combinazioni più "
             "complesse di asset."),
        ],
    )

    # 2.7 Diagnostic plots
    add_heading(doc, "2.7 Diagnostica del modello vincitore (EN tuned, monthly)",
                level=2)
    fig_path = ROOT / "_report_figs" / "sec2_diagnostic_plots.png"
    if fig_path.exists():
        add_figure(
            doc, fig_path,
            caption="Pannelli (dall'alto): (1) returns cumulati target vs "
                    "replica gross e net; (2) drawdowns target vs replica "
                    "net; (3) gross exposure (asse sinistro) e 1-month "
                    "Cornish-Fisher VaR al 99% (asse destro), con tetto "
                    "UCITS 20% segnato in rosso.",
            width_cm=14.5,
        )
    add_para(
        doc,
        "Osservazioni utili dai tre pannelli:",
    )
    add_bullets(
        doc,
        [
            ("Pannello 1 — Returns cumulati: ",
             "la replica segue il target con buona aderenza visiva su "
             "tutto l'OOS; il gap finale fra target (nero) e replica net "
             "(crimson) è ridotto, coerente con TE ≈ 2.8%. La "
             "differenza fra gross (blu) e net (crimson) è "
             "marginalissima — il monthly rebalancing rende i costi "
             "irrilevanti."),
            ("Pannello 2 — Drawdowns: ",
             "i due drawdown si muovono in fase ma con ampiezze "
             "leggermente diverse. La replica tende a sotto-replicare i "
             "drawdown peggiori del target (il MaxDD del target è circa "
             "−29%, la replica si ferma intorno a −16%): è la conseguenza "
             "della gross exposure ridotta. Stessa direzione, ma "
             "ampiezza inferiore — la prima evidenza concreta del Beta "
             "gap."),
            ("Pannello 3 — Gross exposure & VaR: ",
             "la gross exposure oscilla in fascia 0.5–0.8, il VaR resta "
             "ben sotto il tetto UCITS (linea rossa al 20%) per tutto "
             "l'OOS. Il VaR scaling non si è mai attivato in modo "
             "vincolante — segno che la regolarizzazione fa già il "
             "lavoro grosso di limitare l'esposizione."),
        ],
    )
    add_para(
        doc,
        "I top 10 pesi per modulo medio:",
    )
    add_table(
        doc,
        header=["Future", "Peso medio", "Peso finale", "Asset class"],
        rows=[
            ["ES1",  "+0.150", "+0.261", "Equity DM"],
            ["DU1",  "−0.074", "+0.808", "Bond"],
            ["TY1",  "+0.071", "+0.009", "Bond"],
            ["VG1",  "+0.051", "+0.038", "Equity DM"],
            ["NQ1",  "+0.043",  "0.000", "Equity DM"],
            ["GC1",  "+0.039", "+0.015", "Commodity"],
            ["TP1",  "+0.034", "+0.033", "Equity DM"],
            ["TU2",  "+0.020",  "0.000", "Bond"],
            ["RX1",  "+0.015", "+0.084", "Bond"],
            ["CO1",  "+0.014",  "0.000", "Commodity"],
        ],
    )
    add_para(
        doc,
        "ES1 è il driver principale (coerente con la sezione 1: il "
        "future con correlazione più alta col target). I bond hanno pesi "
        "medi piccoli ma il peso finale di DU1 esplode a +0.808 — un "
        "salto di oltre dieci volte rispetto alla media storica. Questa "
        "instabilità della snapshot finale è una flag importante: tutte "
        "le analisi che useranno i pesi correnti (Idee 8, 8b, D, E) "
        "lavoreranno su una configurazione potenzialmente non "
        "rappresentativa.",
    )

    # 2.7 Flags
    flags = [
        ("Baseline = leva eccessiva: ",
         "OLS/Ridge/Lasso/EN default operano a gross exposure ≈ 1.7. Non "
         "sono candidati di consegna; servono solo come reference contro "
         "cui giudicare la regolarizzazione."),
        ("EN tuned monthly (`elnet_monthly`) è il \"modello di "
         "produzione\" presupposto da tutto il notebook a valle: ",
         "Idee 8, 8b, D, D2, E, E2 lo consumano direttamente o via "
         "`best_factory`. Tenere conto che, alla luce delle analisi più "
         "avanzate (Liquidity-aware, NNLS fully-invested), non sarà più "
         "il modello vincitore — sarà necessario eseguire un cambio "
         "centralizzato a valle."),
        ("IR netto strutturalmente negativo (~−0.20): ",
         "il replicatore lineare cattura ~3.9%/anno contro un target di "
         "~4.5%/anno. La differenza è il \"Beta gap\": ce ne occuperemo "
         "esplicitamente più avanti, con NNLS fully-invested e "
         "bootstrap CI dei coefficienti."),
        ("Loss surface piatta vicino all'ottimo (top-5 CV TE entro 1.2%): ",
         "l1_ratio può essere portato a 1.0 senza degradare la TE. Se "
         "futuri pezzi del progetto premieranno la sparsità "
         "(interpretabilità, costi operativi su singoli asset, vincoli "
         "regolamentari per asset class), passare a Lasso puro è una "
         "scelta difendibile."),
        ("Window 208 è il \"common ground\" del notebook: ",
         "verrà condivisa anche da NNLS (cella 26) e Kalman "
         "(init_window). Qualunque modifica al rolling window richiede "
         "un re-run a cascata."),
        ("Snapshot finale dei pesi instabile (DU1 +0.808 finale vs "
         "−0.074 medio): ",
         "tutte le analisi che useranno `weights.iloc[-1]` (Idea 8b "
         "historical replay, Idea E precomputed w_ideal) sono "
         "potenzialmente sensibili a quella settimana. Annotare per la "
         "Sezione sull'historical replay che il look-ahead bias già "
         "dichiarato si combina con questa fragilità."),
        ("VaR scaling mai vincolante per EN tuned: ",
         "il cap UCITS al 20% non è mai stato toccato. Significa che la "
         "regolarizzazione + il rebalance mensile da soli mantengono la "
         "replica dentro i limiti regolamentari. Per modelli più "
         "aggressivi (baselines, multi-moment con λ alto) il cap "
         "diventerà vincolante e influenzerà l'OOS."),
    ]
    add_flag_box(doc, "Flag per il prosieguo del progetto", flags)

    add_para(doc, "")
    return doc


# ---------- SECTION 3: NNLS, Fully-invested NNLS, Kalman ---------------------

def section_3(doc):
    add_heading(doc, "3. Modelli vincolati e modelli dinamici: NNLS, "
                     "Fully-invested NNLS, Kalman", level=1)
    add_para(
        doc,
        "Dopo i baseline lineari e l'Elastic Net tunato la Sezione 2 ha "
        "lasciato due flag aperte: gross exposure ancora ben sotto 1 (Beta "
        "gap) e IR netto strutturalmente negativo. Questa sezione affronta "
        "entrambi i temi introducendo tre nuovi modelli — due varianti "
        "NNLS che agiscono sui vincoli di portafoglio, e un Kalman filter "
        "che agisce sulla dinamica dei pesi.",
    )

    # 3.1 NNLS
    add_heading(doc, "3.1 Non-negative Lasso (long-only)", level=2)
    add_para(
        doc,
        "Il primo intervento è strutturale: vietare le posizioni corte. "
        "Tre motivazioni si compongono:",
    )
    add_bullets(
        doc,
        [
            ("Regolamentare: ",
             "molti fondi UCITS proibiscono esplicitamente le posizioni "
             "corte in derivati salvo che per copertura. Un replicatore "
             "long-only è quindi il prodotto regolamentarmente più "
             "semplice da distribuire."),
            ("Statistica: ",
             "il vincolo w ≥ 0 agisce come un prior molto forte. Esclude "
             "le configurazioni più volatili e instabili dei coefficienti "
             "e funziona da regolarizzatore implicito, producendo pesi "
             "tipicamente più sparsi e stabili."),
            ("Empirica: ",
             "il target Monster Index è strutturalmente long (equity, "
             "bond, hedge fund a beta positivo). Permettere pesi negativi "
             "sui futures equivale a dare all'ottimizzatore una libertà "
             "che spesso usa per inseguire rumore."),
        ],
    )
    add_para(
        doc,
        "Implementativamente è un Lasso con vincolo di non-negatività. "
        "Centriamo il target settimanale (yc = y − ȳ) per assorbire un "
        "intercetta implicita, poi risolviamo via scipy.optimize.nnls "
        "(algoritmo di Lawson-Hanson) la:",
    )
    add_formula(doc, "min_w  ‖Xw − yc‖²    s.t.   wj ≥ 0  ∀j")
    add_para(
        doc,
        "Post-processing: pesi con valore assoluto sotto la soglia "
        "0.005 vengono azzerati (sparsity_thresh), il resto resta "
        "invariato. Il modello viene inserito nel backtest() generico — "
        "stesso engine walk-forward dei baseline (Protocollo A della "
        "Sezione 2.5), con rolling window 208 settimane e rebalance ogni "
        "4 settimane.",
    )

    fig = ROOT / "_report_figs" / "sec3_nnls_weights.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Pesi medi NNLS long-only su tutto l'OOS. Distribuzione "
                    "fortemente concentrata su ES1, NQ1, VG1, TP1 (equity "
                    "DM) con un contributo importante di GC1 (oro).",
            width_cm=14.5,
        )
    add_para(
        doc,
        "Risultati sul comune OOS (496 settimane, dal 25 ottobre 2011):",
    )
    add_table(
        doc,
        header=["Modello", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett."],
        rows=[
            ["EN tuned, monthly", "0.0281", "−0.199", "0.853", "0.64",
             "0.0182"],
            ["NNLS long-only",    "0.0275", "−0.169", "0.860", "0.64",
             "0.0160"],
        ],
    )
    add_para(
        doc,
        "Il vincolo di non-negatività paga su tutta la riga: TE leggermente "
        "migliore (−6 bps in termini di volatilità di tracking), IR "
        "migliore di 3 punti, correlazione più alta, turnover più basso. "
        "La gross exposure resta però vicino a 0.64 — il Beta gap non si "
        "chiude, perché senza vincolo di leva il solver \"prende meno\" "
        "che con una decisione esplicita di essere fully-invested.",
    )

    # 3.2 Fully-invested NNLS
    add_heading(doc, "3.2 Fully-invested NNLS — chiusura del Beta gap",
                level=2)
    add_para(
        doc,
        "Il passo successivo è aggiungere il vincolo di pieno impiego del "
        "capitale. La gross exposure diventa esattamente 1 per costruzione "
        "(Beta ≈ 1 \"by design\"), togliendo al modello la possibilità "
        "implicita di sotto-investire. Il problema diventa:",
    )
    add_formula(
        doc,
        "min_w ‖y − Xw‖²    s.t.   wj ≥ 0  ∀j,   Σⱼ wj = 1"
    )
    add_para(
        doc,
        "Si tratta di una programmazione quadratica con vincoli misti di "
        "bound (caja [0, 1]) e una equality lineare (somma = 1). La "
        "risolviamo con scipy.optimize.minimize, metodo SLSQP, warm-start "
        "feasible (w₀ = 1/k), tolleranza 10⁻¹⁰. Non c'è intercetta: il "
        "vincolo somma = 1 fissa già il livello del portafoglio sul "
        "livello dei futures, che è naturalmente vicino a quello del "
        "target. Sparsity threshold sempre 0.005, con rinormalizzazione "
        "post-soglia per mantenere Σw = 1.",
    )
    add_para(
        doc,
        "Sanity check sulla soluzione media (post VaR scaling):",
    )
    add_bullets(
        doc,
        [
            ("Σw ≈ 1.000, ", "vincolo rispettato."),
            ("Gross exposure ≈ 1.000, ", "Beta ≈ 1 by construction."),
            ("Peso massimo medio 0.271 (ES1), ", "nessuna concentrazione "
             "esagerata: tutta la massa è distribuita su più asset."),
            ("Active assets ≈ 8.2 su 11, ", "sparsità moderata: solo tre "
             "futures restano sistematicamente sotto soglia."),
        ],
    )
    fig = ROOT / "_report_figs" / "sec3_fi_nnls_weights.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Composizione media del portafoglio fully-invested NNLS. "
                    "ES1 al 27%, distribuzione bilanciata fra equity DM "
                    "(~32% totale), bond US/EU (~26% totale), commodities "
                    "(~6%).",
            width_cm=14.5,
        )
    add_para(
        doc,
        "Risultati a confronto con NNLS standard sull'OOS comune:",
    )
    add_table(
        doc,
        header=["Modello", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett."],
        rows=[
            ["NNLS long-only (gross < 1)", "0.0275", "−0.169", "0.860",
             "0.64", "0.0160"],
            ["Fully-invested NNLS (Σw=1)", "0.0275", "−0.023", "0.860",
             "1.00", "0.0162"],
        ],
    )
    add_para(
        doc,
        "Il salto di IR è la firma del closing del Beta gap: IR netto "
        "passa da −0.169 a −0.023, di fatto azzerato. La TE è invariata "
        "(0.0275): il vincolo non costa accuratezza di tracking. Il "
        "turnover sale appena di un decimo di punto base, e la gross "
        "exposure si attesta esattamente a 1, dentro il limite UCITS. È "
        "la prima volta in tutto il notebook che vediamo un modello con "
        "IR netto sostanzialmente non distinguibile da zero — la replica "
        "cattura il drift del target.",
    )
    add_para(
        doc,
        "La composizione media (≈ 32% equity, ≈ 26% bond, ≈ 6% "
        "commodity) lascia un \"missing 36%\" rispetto al sample medio "
        "che il modello compensa via ribilanciamento interno. Quel gap "
        "è quanto resta dell'alpha hedge fund non replicabile coi futures "
        "liquidi — circa 50 bps/anno di drift residuo, lo stesso che "
        "ricorre nell'IR di tutti gli altri modelli.",
    )

    # 3.3 Re-baseline
    add_heading(doc, "3.3 Re-baseline dei modelli a window 208", level=2)
    add_para(
        doc,
        "Il grid search della Sezione 2 ha scelto window = 208 (4 anni) "
        "per l'Elastic Net tunato, NNLS e Fully-invested NNLS hanno "
        "riusato la stessa scelta, e il Kalman filter — che vedremo nella "
        "prossima sotto-sezione — userà 208 settimane di warm-up. Per "
        "evitare lo slicing-trick del Protocollo C della Sezione 2.5 in "
        "tutte le tabelle successive, qui re-runniamo i quattro baseline "
        "(OLS, Ridge, Lasso, EN default) con window = 208 invece che 156. "
        "Da questo punto in avanti tutti i modelli del notebook "
        "condividono la stessa finestra di training e quindi lo stesso "
        "intervallo OOS (496 settimane, OOS start 25 ottobre 2011).",
    )
    add_para(
        doc,
        "L'effetto del passaggio 156 → 208 settimane sui baseline è "
        "uniforme e modesto:",
    )
    add_table(
        doc,
        header=["Modello", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett."],
        rows=[
            ["OLS",        "0.0274", "−0.179", "0.861", "1.56", "0.069"],
            ["Ridge",      "0.0274", "−0.181", "0.860", "1.54", "0.068"],
            ["Lasso",      "0.0274", "−0.177", "0.861", "1.51", "0.067"],
            ["EN default", "0.0274", "−0.178", "0.861", "1.54", "0.068"],
        ],
    )
    add_para(
        doc,
        "Più dati per il training (208 vs 156) producono coefficienti "
        "più stabili: gross exposure scende da ≈ 1.69 a ≈ 1.54, "
        "turnover settimanale cala del 23%, TE marginalmente migliore. "
        "Il livello resta comunque non-UCITS-compliant (gross > 1.0). "
        "I baseline restano quindi reference, non candidati di consegna.",
    )

    # 3.4 Kalman
    add_heading(doc, "3.4 Kalman filter — pesi dinamici online", level=2)
    add_para(
        doc,
        "Tutti i modelli visti finora condividono lo stesso schema: "
        "training su finestra rolling, refit periodico, pesi statici fra "
        "un refit e il successivo. Il Kalman filter inverte la logica: "
        "mantiene una stima dei pesi continuamente aggiornata, "
        "incorporando ogni nuova osservazione del target senza scartare "
        "quelle vecchie. È particolarmente adatto a problemi con \"regime "
        "drift\" come questo, dove l'esposizione del target evolve nel "
        "tempo (gli hedge fund spostano asset class fra calm e crisi).",
    )
    add_heading(doc, "Formulazione state-space", level=3)
    add_para(
        doc,
        "Il modello tratta i pesi del portafoglio come variabile latente "
        "βt che evolve secondo un random walk; la \"misurazione\" è il "
        "return settimanale del target.",
    )
    add_formula(
        doc,
        "Equazione di stato:   βt = β_{t−1} + wt,    "
        "wt ~ N(0, Q)"
    )
    add_formula(
        doc,
        "Equazione di osservazione:   yt = Xtᵀ βt + vt,    "
        "vt ~ N(0, R)"
    )
    add_para(
        doc,
        "Due parametri governano il comportamento del filtro:",
    )
    add_bullets(
        doc,
        [
            ("Q (process noise covariance): ",
             "diagonale Q = q · I, controlla l'agilità del filtro. q "
             "grande → pesi che possono cambiare velocemente, segnale "
             "dominante; q piccolo → pesi quasi statici, smoothing "
             "dominante."),
            ("R (observation noise variance): ",
             "quanto del rumore di mercato attribuiamo a vt, cioè "
             "quanta varianza del target è da \"ignorare\" come rumore. "
             "Inizializzato sulla varianza dei residui Ridge del warm-up; "
             "non viene poi più aggiornato."),
        ],
    )
    add_heading(doc, "Inizializzazione e aggiornamento online", level=3)
    add_para(
        doc,
        "Per evitare che il filtro parta cieco, β₀ viene inizializzato "
        "con i coefficienti di una Ridge (α = 1) regredita sulle prime "
        "208 settimane. La covarianza iniziale è P₀ = 0.01 · I. Da lì in "
        "poi, ogni settimana t ≥ 208:",
    )
    add_bullets(
        doc,
        [
            ("Predict: ",
             "P ← P + Q. Espandere l'incertezza per tenere conto del "
             "random walk dei pesi."),
            ("VaR scaling: ",
             "calcolare il VaR Cornish-Fisher sull'esposizione corrente "
             "e ridurre β di scale = min(1, VAR_MAX / VaR)."),
            ("Out-of-sample step: ",
             "next_ret = Xtᵀ · β_scaled. È la previsione del filtro per "
             "la settimana t, registrata come OOS prima di aggiornare "
             "con yt."),
            ("Update: ",
             "innovazione e = yt − Xtᵀβ; gain K = P·Xt / (XtᵀP·Xt + R); "
             "β ← β + K·e; P ← P − K·Xtᵀ·P."),
        ],
    )
    add_para(
        doc,
        "Differenza fondamentale rispetto al walk-forward classico: il "
        "Kalman non rifa un fit, aggiorna lo stato. Tutta la storia "
        "passata contribuisce a β tramite la sequenza di update, ma con "
        "memoria controllata dal rapporto R/Q. Il VaR scaling resta "
        "identico al Protocollo A della Sezione 2.5 (rolling 156 "
        "settimane di replica returns).",
    )
    add_heading(doc, "Grid search su q e selezione del filtro vincente",
                level=3)
    add_para(
        doc,
        "Q viene scelto con una piccola ricerca su scala log: "
        "q ∈ {10⁻⁷, 10⁻⁶, 10⁻⁵, 10⁻⁴, 10⁻³}. Il punteggio è la TE netta "
        "su tutto l'OOS (496 settimane). R viene stimato in modo "
        "data-driven dai residui Ridge del warm-up e non viene incluso "
        "nella ricerca.",
    )
    add_table(
        doc,
        header=["q", "TE net", "IR net", "Gross exp.", "Turnover/sett."],
        rows=[
            ["10⁻⁷", "0.0276", "−0.154", "0.48", "0.0077"],
            ["10⁻⁶", "0.0279", "−0.142", "0.48", "0.0082"],
            ["10⁻⁵", "0.0288", "−0.088", "0.50", "0.0112"],
            ["10⁻⁴", "0.0297", "−0.062", "0.58", "0.0239"],
            ["10⁻³", "0.0308", "−0.191", "0.76", "0.0624"],
        ],
    )
    add_para(
        doc,
        "Il vincitore è q = 10⁻⁷, il valore più piccolo della griglia. Il "
        "filtro \"preferisce\" pesi quasi statici. Tre letture si "
        "compongono in una conclusione robusta:",
    )
    add_bullets(
        doc,
        [
            ("Stabilità strutturale del target: ",
             "il fatto che il q ottimale sia il minimo (cioè dare al "
             "modello la minima libertà di muovere i pesi) significa che "
             "la composizione strategica del Monster Index è stabile sul "
             "campione 2011–2021. Le variazioni di esposizione che "
             "esistono sono incorporate naturalmente nel termine di "
             "innovazione settimanale, senza bisogno di un q grande."),
            ("Costi sotto controllo: ",
             "turnover settimanale 0.0077 contro 0.0162 di Fully-invested "
             "NNLS e 0.018 di EN tunato monthly. È il modello con "
             "turnover più basso, e l'ordine di magnitudine resta valido "
             "anche per q = 10⁻⁶."),
            ("Sensibilità a Q: ",
             "passando da 10⁻⁷ a 10⁻³ il turnover si moltiplica per "
             "otto, la gross exposure cresce da 0.48 a 0.76, e l'IR — "
             "dopo un minimo locale a q = 10⁻⁴ (−0.062) — peggiora "
             "drasticamente. Conferma che un q troppo alto trasforma il "
             "filtro in un inseguitore di rumore."),
        ],
    )

    box = open_methodology_box(
        doc, "Nota statistica — scelta di q senza eval set held-out"
    )
    box_para(
        box,
        "A differenza dell'Elastic Net, q non è scelto via "
        "TimeSeriesSplit su un search set separato: il punteggio "
        "(TE netta) è calcolato sullo stesso OOS che poi viene "
        "riportato come performance del Kalman. C'è quindi una "
        "componente di in-sample optimisation. Il bias è limitato per "
        "tre motivi:",
    )
    box_bullets(
        box,
        [
            ("Griglia minimale: ",
             "5 valori di q in scala log su un solo parametro scalare. "
             "Probabilità di overfit minima."),
            ("Optimum agli estremi: ",
             "il vincitore (10⁻⁷) è il minimo della griglia. "
             "Estendere ancora più in basso porterebbe verso un filtro "
             "completamente statico, con risultati che convergerebbero a "
             "Ridge sulla finestra di warm-up. La scelta è dunque \"il "
             "filtro più stabile possibile\", senza spazio per "
             "sensibilità fine."),
            ("Conferma indipendente: ",
             "le sezioni successive sulla bootstrap CI dei coefficienti "
             "mostrano che i pesi del Kalman sono indistinguibili da "
             "quelli del fit medio della Ridge nel sample. È evidenza "
             "che q = 10⁻⁷ non è un over-fit ma un sottostato del "
             "modello."),
        ],
    )

    # 3.5 Side-by-side
    add_heading(doc, "3.5 Confronto side-by-side dei replicatori",
                level=2)
    add_para(
        doc,
        "Quadro sintetico finale dei modelli analizzati fin qui, tutti "
        "sull'OOS comune (496 settimane, 25 ottobre 2011 → 20 aprile "
        "2021):",
    )
    add_table(
        doc,
        header=["Modello", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett."],
        rows=[
            ["OLS (208)",                 "0.0274", "−0.179", "0.861", "1.56", "0.069"],
            ["Ridge (208)",               "0.0274", "−0.181", "0.860", "1.54", "0.068"],
            ["Lasso (208)",               "0.0274", "−0.177", "0.861", "1.51", "0.067"],
            ["EN default (208)",          "0.0274", "−0.178", "0.861", "1.54", "0.068"],
            ["EN tuned, monthly",         "0.0281", "−0.199", "0.853", "0.64", "0.0182"],
            ["NNLS long-only",            "0.0275", "−0.169", "0.860", "0.64", "0.0160"],
            ["Fully-invested NNLS",       "0.0275", "−0.023", "0.860", "1.00", "0.0162"],
            ["Kalman (q = 10⁻⁷)",         "0.0276", "−0.154", "0.859", "0.48", "0.0077"],
        ],
    )
    fig = ROOT / "_report_figs" / "sec3_side_by_side.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Returns cumulati netti delle tre famiglie di "
                    "replicatori (EN tuned, NNLS, Kalman) vs target. Le "
                    "tre curve sono quasi indistinguibili nel tratto "
                    "centrale; le differenze sono nei drawdown e nel "
                    "tracking durante le crisi.",
            width_cm=14.5,
        )
    add_para(
        doc,
        "Quattro osservazioni operative:",
    )
    add_bullets(
        doc,
        [
            ("Plateau della Tracking Error: ",
             "tutti i modelli convergono attorno a TE = 2.74–2.81%. La "
             "differenza fra il migliore e il peggiore è di 7 bps di "
             "deviazione standard di tracking — irrilevante in pratica. "
             "Significa che il problema della replica con regressione "
             "lineare ha un floor di TE imposto dalla parte "
             "irriducibile della varianza dovuta al return smoothing "
             "di HFRX (vedi flag della Sezione 1)."),
            ("La gara si decide su gross exposure, turnover e IR: ",
             "tre metriche correlate fra loro. Fully-invested NNLS "
             "vince sul Beta gap (IR ≈ 0), Kalman vince sui costi "
             "(turnover 0.0077, gross 0.48), EN tunato e NNLS sono "
             "intermedi su entrambi i fronti."),
            ("Il \"floor\" negativo di IR ≈ −0.15: ",
             "tutti i modelli tranne il Fully-invested NNLS hanno IR "
             "netto vicino a −0.15. Quel valore rappresenta il drift "
             "non replicabile (~50 bps/anno di alpha hedge fund "
             "puro). Senza un vincolo che obblighi all'investimento "
             "pieno, qualsiasi regressione lo lascia sul tavolo."),
            ("Baselines a 208 settimane restano fuori UCITS: ",
             "anche con la finestra più lunga, gross exposure sopra "
             "1.5. Non sono utilizzabili come prodotto distribuito; "
             "restano riferimenti metrici."),
        ],
    )

    # 3.6 Flags
    flags = [
        ("Da questa sezione in poi l'OOS è unico (496 settimane): ",
         "tutti i modelli con window = 208 (baseline re-runnati, EN "
         "tunato, NNLS, Fully-invested NNLS, Kalman) condividono il "
         "medesimo intervallo OOS. Le tabelle successive non "
         "richiederanno più il re-slicing del Protocollo C."),
        ("Fully-invested NNLS è il primo candidato serio di consegna: ",
         "IR ≈ 0, gross = 1 (UCITS-compliant), turnover paragonabile a "
         "EN tunato. Va incluso esplicitamente in tutte le analisi a "
         "valle (bootstrap CI, stress test, conformal). Tenere d'occhio: "
         "in alcuni run successivi il notebook continua a usare "
         "elnet_monthly come \"modello di produzione\" — è un'eredità "
         "del primo grid search che andrà sostituita centralmente."),
        ("Kalman è il candidato per il fronte costi: ",
         "se la priorità futura sarà ridurre il rebalancing (ad esempio "
         "in presenza di costi più alti, o per asset poco liquidi), "
         "Kalman parte già con turnover 0.0077. Le estensioni "
         "production-oriented (liquidity-aware fit, adaptive band) "
         "useranno questi numeri come baseline da migliorare."),
        ("q = 10⁻⁷ è il setting Kalman per il resto del notebook: ",
         "il valore vincente viene salvato in best_q ed usato anche "
         "dall'adaptive band stateful (versione corretta del wrapper) "
         "che vedremo più avanti. Non è più un parametro libero."),
        ("Il \"floor\" di IR ≈ −0.15 ha un significato strutturale: ",
         "rappresenta la frazione di drift del target che nessun "
         "modello lineare può catturare. Le analisi di stress / "
         "conformal devono interpretare i CI di IR alla luce di "
         "questo: se il CI di un modello include solo valori >= −0.15, "
         "il modello sta catturando tutto il replicabile."),
        ("Sparsity threshold (0.005) per entrambi gli NNLS: ",
         "fissato a mano sulla base di sanity check di stabilità dei "
         "pesi. Non è ottimizzato. Potrebbe essere un parametro da "
         "tunare in versioni future del progetto, soprattutto se si "
         "vuole un portafoglio con un numero ridotto di asset attivi."),
        ("NNLS standard e Fully-invested NNLS condividono la stessa "
         "sparsità media (8 asset attivi su 11): ",
         "i bond a breve scadenza (TU2, DU1) e talvolta CO1 finiscono "
         "sotto soglia. Significato pratico: la parte di hedging del "
         "portafoglio è concentrata su scadenze lunghe (RX1, TY1). "
         "Da ricordare quando giudicheremo modelli che usano feature "
         "esplicite di asset class (Idea 4 constrained, multi-moment)."),
    ]
    add_flag_box(doc, "Flag per il prosieguo del progetto", flags)

    add_para(doc, "")
    return doc


# ---------- SECTION 4: Bootstrap CI + Span analysis --------------------------

def section_4(doc):
    add_heading(doc, "4. Intervalli di confidenza bootstrap e analisi dello span",
                level=1)
    add_para(
        doc,
        "Le sezioni precedenti hanno classificato i modelli usando "
        "punti-stima delle metriche (TE, IR, correlazione, gross "
        "exposure, turnover). Questa sezione fa una cosa diversa: "
        "associa a ogni metrica un intervallo di confidenza al 90% "
        "tramite stationary bootstrap, e affianca un'analisi geometrica "
        "del problema (lo \"span analysis\") che fissa il limite teorico "
        "di quanto un modello lineare possa replicare di una serie "
        "obiettivo data l'universo dei regressori. Insieme, le due analisi "
        "rispondono a una domanda che le tabelle puntuali non affrontano: "
        "le differenze di performance fra i nostri modelli sono "
        "statisticamente significative, o sono compatibili con il "
        "rumore campionario?",
    )

    # 4.1 Why bootstrap CI
    add_heading(doc, "4.1 Perché servono gli intervalli di confidenza",
                level=2)
    add_para(
        doc,
        "Una point estimate di TE = 2.75% e una di TE = 2.81% sembrano "
        "diverse, ma se il rumore campionario sulla TE è dell'ordine "
        "dell'1.5%, le due stime sono di fatto indistinguibili. Senza "
        "una banda di incertezza il ranking dei modelli rischia di "
        "essere artifizioso. Quattro metriche entrano nel bootstrap, "
        "scelte perché catturano dimensioni indipendenti della qualità "
        "della replica:",
    )
    add_bullets(
        doc,
        [
            ("Tracking Error: ", "deviazione standard del tracking "
             "difference. Quanto la replica avvolge da vicino la "
             "traiettoria del target."),
            ("Information Ratio: ", "rendimento medio del tracking "
             "difference diviso TE. Misura il bias dopo i costi; "
             "idealmente zero."),
            ("Correlation: ", "co-movimento direzionale, indipendente "
             "dalla scala. Una replica con Beta basso ma alta "
             "correlazione \"si muove insieme\" al target ma a "
             "magnitudo ridotta."),
            ("Beta: ", "pendenza della regressione replica su target "
             "(cov / var(target)). Misura la scala. Beta = 1 significa "
             "che la replica è dimensionata correttamente; sotto 1 è "
             "sotto-investita, sopra 1 è sovra-leveraged. È la metrica "
             "che spiega perché una replica con gross exposure bassa "
             "può avere ottima correlazione ma sotto-tracciare il drift."),
        ],
    )

    # 4.2 Stationary bootstrap
    add_heading(doc, "4.2 Stationary block bootstrap (Politis-Romano)",
                level=2)
    add_para(
        doc,
        "Un bootstrap classico (riassortimento i.i.d. delle settimane) "
        "distruggerebbe la struttura di autocorrelazione dei returns — "
        "volatility clustering, drifts lenti, code in stress. I CI "
        "risultanti sarebbero falsamente stretti. Lo stationary "
        "bootstrap di Politis-Romano risolve il problema campionando "
        "blocchi consecutivi di lunghezza casuale (geometrica, media 8 "
        "settimane ≈ 2 mesi), preservando la struttura seriale locale "
        "ma ottenendo blocchi quasi i.i.d. fra una replica e l'altra.",
    )
    add_para(
        doc,
        "Procedura operativa, per ciascuna delle 500 ripetizioni "
        "bootstrap:",
    )
    add_bullets(
        doc,
        [
            ("Costruzione di un campione: ",
             "si parte da un indice random, si aggiungono settimane "
             "consecutive; con probabilità p = 1/8 a ogni passo si "
             "salta a un nuovo indice random (nuovo blocco). Si "
             "continua fino a riempire 496 settimane (la lunghezza "
             "dell'OOS)."),
            ("Calcolo delle metriche sul campione: ",
             "TE, IR, correlazione, Beta — usando i returns della "
             "replica e del target letti dagli indici campionati."),
        ],
    )
    add_para(
        doc,
        "Alla fine si dispone di 500 valori per ogni metrica e ogni "
        "modello; la media è la stima centrale, e i quantili 5° e 95° "
        "definiscono il CI al 90%. Eseguiamo il bootstrap su tutti gli "
        "otto modelli con window = 208 introdotti nelle Sezioni 2 e 3.",
    )

    # 4.3 Results
    add_heading(doc, "4.3 Risultati del bootstrap", level=2)
    add_para(
        doc,
        "Tabella sintetica: per ogni modello, media e intervallo 5–95% "
        "delle quattro metriche (500 repliche bootstrap, blocchi medi "
        "di 8 settimane, seed fisso = 7 per riproducibilità).",
    )
    add_table(
        doc,
        header=["Modello", "TE mean [5%, 95%]", "IR mean [5%, 95%]",
                "Corr mean [5%, 95%]", "Beta mean [5%, 95%]"],
        rows=[
            ["OLS",                  "0.027 [0.022, 0.035]",
             "−0.21 [−0.69, +0.24]", "0.862 [0.834, 0.891]",
             "0.780 [0.723, 0.849]"],
            ["Ridge",                "0.027 [0.022, 0.035]",
             "−0.21 [−0.70, +0.24]", "0.861 [0.832, 0.892]",
             "0.776 [0.716, 0.846]"],
            ["Lasso",                "0.027 [0.022, 0.035]",
             "−0.20 [−0.69, +0.24]", "0.862 [0.833, 0.891]",
             "0.778 [0.720, 0.847]"],
            ["EN default",           "0.027 [0.022, 0.035]",
             "−0.21 [−0.69, +0.24]", "0.862 [0.833, 0.891]",
             "0.778 [0.721, 0.848]"],
            ["EN tuned",             "0.028 [0.021, 0.037]",
             "−0.23 [−0.73, +0.25]", "0.855 [0.820, 0.890]",
             "0.702 [0.640, 0.773]"],
            ["NNLS long-only",       "0.027 [0.022, 0.035]",
             "−0.19 [−0.64, +0.24]", "0.861 [0.830, 0.893]",
             "0.782 [0.722, 0.852]"],
            ["Kalman (q = 10⁻⁷)",    "0.027 [0.022, 0.035]",
             "−0.18 [−0.68, +0.31]", "0.859 [0.832, 0.889]",
             "0.710 [0.647, 0.784]"],
            ["Fully-invested NNLS",  "0.027 [0.022, 0.035]",
             "−0.04 [−0.47, +0.38]", "0.861 [0.829, 0.892]",
             "0.782 [0.724, 0.852]"],
        ],
    )

    add_heading(doc, "Lettura per metrica", level=3)
    add_para(
        doc,
        "TE — sette modelli su otto sono indistinguibili. Il CI è "
        "[0.022, 0.035] per chiunque (ampiezza 1.3 punti percentuali); "
        "le medie variano fra 0.027 e 0.028. EN tuned è marginalmente "
        "più alto (0.028 di media, 95° percentile a 0.037), ma il suo "
        "CI si sovrappone completamente con quello degli altri. "
        "Conclusione: nessun modello è statisticamente migliore di un "
        "altro sulla TE.",
    )
    add_para(
        doc,
        "IR — tutti gli IR medi sono negativi (range −0.04 ÷ −0.23), e "
        "ogni intervallo include zero. Anche il Fully-invested NNLS, "
        "con IR medio −0.045 (la migliore stima puntuale), ha CI "
        "[−0.47, +0.38]: il punto zero è ben dentro la banda. "
        "Conclusione formale: con 90% di confidenza non possiamo "
        "rifiutare l'ipotesi nulla \"IR = 0\" per nessuno dei modelli. "
        "I ranking di IR fatti finora sono ranking di stime puntuali, "
        "non di parametri sottostanti.",
    )
    add_para(
        doc,
        "Correlation — tutti i CI si sovrappongono in [0.82, 0.89]. "
        "EN tuned ha il lower bound più basso (0.820), gli altri "
        "stanno fra 0.830 e 0.834. Differenze al margine, non sufficienti "
        "a fondare un ranking statistico.",
    )
    add_para(
        doc,
        "Beta — qui finalmente emerge una distinzione netta. Tutti i "
        "modelli hanno Beta media fra 0.70 e 0.78, e ogni 95° "
        "percentile è strettamente inferiore a 1: il CI 90% non tocca "
        "mai 1.0. La separazione fra modelli c'è: EN tuned e Kalman "
        "stanno significativamente sotto (medie 0.70 e 0.71, 95° "
        "percentile sotto 0.78); OLS, Ridge, Lasso, EN default, NNLS, "
        "Fully-invested NNLS stanno tutti vicini a 0.78. La differenza "
        "fra i due gruppi (~7 punti di Beta) eccede l'ampiezza media "
        "del CI ed è quindi statisticamente significativa.",
    )

    # 4.4 Sotto-investiment puzzle (Box)
    box = open_methodology_box(
        doc,
        "Perché anche Fully-invested NNLS ha Beta ≈ 0.78"
    )
    box_para(
        box,
        "A prima vista è un paradosso. Il Fully-invested NNLS ha gross "
        "exposure = 1 per costruzione, quindi \"sa\" di dover essere "
        "completamente investito. Eppure la sua Beta bootstrappata "
        "resta a 0.78. La spiegazione è geometrica:",
    )
    box_para(
        box,
        "    Beta = Cov(replica, target) / Var(target) "
        "= Corr(replica, target) · σ_replica / σ_target",
        italic=True,
    )
    box_para(
        box,
        "Con Corr ≈ 0.86 e σ_replica / σ_target ≈ 0.0485 / 0.0537 ≈ 0.90 "
        "si ottiene Beta ≈ 0.86 · 0.90 ≈ 0.78. Il vincolo Σw = 1 "
        "fissa il livello del portafoglio (e infatti chiude il gap "
        "dell'IR, ora −0.04 invece di −0.17), ma non altera la "
        "*varianza* della replica, che è limitata dalla varianza "
        "raggiungibile nel sottospazio generato dai futures. Il gap "
        "di Beta = 1 − 0.78 ≈ 0.22 non è un'inefficienza del "
        "modellatore: è la quota di varianza del target che esce dal "
        "linear span dei nostri 11 regressori. L'analisi dello span "
        "(prossima sotto-sezione) lo quantifica esattamente.",
    )

    # 4.5 Span analysis
    add_heading(doc, "4.4 Analisi dello span: il limite teorico della replica",
                level=2)
    add_para(
        doc,
        "Per ogni serie target (le tre componenti del Monster Index e "
        "il Monster Index stesso) facciamo una OLS sul sample completo "
        "contro l'intero universo degli 11 futures, oppure contro un "
        "sotto-insieme per asset class. L'R² di queste regressioni dà "
        "una stima dall'alto della frazione di varianza che un modello "
        "lineare sui futures può raggiungere — è il massimo teorico, "
        "non una performance effettiva (perché usa tutto il sample e "
        "tutti i regressori, senza penalità).",
    )
    add_table(
        doc,
        header=["Target", "Tutti gli 11 futures", "Solo equity (5)",
                "Solo bond (4)", "Solo commodity (2)"],
        rows=[
            ["HFRXGL",            "0.275", "0.236", "0.041", "0.080"],
            ["MXWO",              "0.963", "0.952", "0.161", "0.230"],
            ["LEGATRUU",          "0.512", "0.046", "0.340", "0.258"],
            ["Monster Index",     "0.786", "0.732", "0.073", "0.212"],
        ],
    )

    add_para(
        doc,
        "Quattro letture ne escono con forza:",
    )
    add_bullets(
        doc,
        [
            ("MXWO è quasi perfettamente replicabile: ",
             "R² = 0.963 sui 5 equity futures. Inevitabile — i futures "
             "ES1, NQ1, VG1, TP1, LLL1 spannano già un'ampia parte "
             "dell'equity globale."),
            ("HFRXGL è quasi inafferrabile: ",
             "R² = 0.275. Solo il 27.5% della varianza di HFRX è nel "
             "linear span dei nostri 11 futures. Il restante 72.5% è "
             "varianza idiosincratica degli hedge fund manager — "
             "posizioni illiquide, strategie discretionary, exposures "
             "esotiche che semplicemente non esistono nel nostro "
             "universo di asset."),
            ("LEGATRUU è solo a metà replicabile (R² = 0.512), e "
             "principalmente *non* dai bond futures: ",
             "i 4 bond futures spiegano solo il 34% di LEGATRUU. La "
             "ragione: LEGATRUU è un aggregato globale (US, EU, JP, EM, "
             "high yield, currency-hedged) mentre i nostri futures "
             "sono solo Treasury US e Bund tedeschi a 2y/10y. Manca "
             "tutta l'esposizione cross-currency, credito, EM, breakeven "
             "inflation."),
            ("Monster Index: R² teorico massimo = 0.786. ",
             "È il limite superiore: nessun modello lineare sui 11 "
             "futures può catturare più del 78.6% della varianza del "
             "Monster. Questo si traduce direttamente nei conti del "
             "Beta: √0.786 ≈ 0.886. Una replica linear-perfetta "
             "potrebbe raggiungere Beta al massimo 0.886 (e questo "
             "solo in-sample, senza regolarizzazione, senza VaR cap, "
             "senza costi). Le nostre repliche OOS si fermano a "
             "≈ 0.78 — circa 10 punti sotto il massimo teorico, "
             "perdita ragionevolmente attribuibile a OOS + regolarizzazione "
             "+ scaling del rischio."),
        ],
    )

    fig = ROOT / "_report_figs" / "sec4_span_1.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Mappa di correlazione fra ciascun future e le tre "
                    "componenti del target + il Monster Index. Conferma "
                    "graficamente i risultati dell'R²: equity futures "
                    "fortemente correlati con MXWO, bond futures più "
                    "tiepidi con LEGATRUU, HFRX correlato \"in modo "
                    "diluito\" con tutti.",
            width_cm=13,
        )
    fig = ROOT / "_report_figs" / "sec4_span_2.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Returns cumulati delle tre componenti del Monster "
                    "Index e del target. Il Monster (nero) cammina come "
                    "una media pesata delle tre, dominato dall'andamento "
                    "smussato di HFRX (crimson) — la componente con il "
                    "drift più basso e la varianza più strana.",
            width_cm=15,
        )

    add_para(
        doc,
        "Lo span analysis chiude logicamente il discorso aperto dal "
        "bootstrap: il \"Beta gap\" universale non è un fallimento dei "
        "modelli, è una proprietà del problema. Per superare il muro "
        "0.78 di Beta servirebbero almeno una delle seguenti modifiche:",
    )
    add_bullets(
        doc,
        [
            ("Universo più ampio: ",
             "aggiungere futures su credito (CDS), volatilità (VIX), "
             "currency carry. L'R² su HFRX salirebbe rispetto al "
             "0.275 attuale."),
            ("Modelli non lineari: ",
             "kernel methods, random forest, MLP. Catturano "
             "non-linearità (regime switching, vol clustering) che il "
             "lineare ignora. Costo: minore interpretabilità, maggior "
             "rischio di overfitting."),
            ("Modelli generativi: ",
             "fitting di una distribuzione condizionale di HFRX dati i "
             "futures, e ribilanciamento basato su scenari simulati. "
             "Lavoro di tesi a sé."),
        ],
    )

    # 4.5 Final verdict
    add_heading(doc, "4.5 Verdetto e implicazioni per il prosieguo", level=2)
    add_para(
        doc,
        "I risultati del bootstrap obbligano a riformulare alcune "
        "conclusioni che le sezioni precedenti avevano tratto sulle "
        "stime puntuali. In modo netto:",
    )
    add_bullets(
        doc,
        [
            ("Su TE e IR, statisticamente, non c'è un vincitore. ",
             "La differenza fra il \"peggior modello\" (EN tuned, TE "
             "0.0277) e il \"migliore\" (OLS, TE 0.0270) è ben dentro "
             "il rumore campionario. Stessa cosa per IR: tutti i CI "
             "includono zero. Conclusione formale: la scelta del "
             "modello di consegna non può essere giustificata sulla "
             "base di TE o IR puntuali — vanno usati altri criteri."),
            ("EN tuned è l'unico modello statisticamente dominato. ",
             "Beta media 0.702, CI [0.640, 0.773] — significativamente "
             "sotto a OLS/Ridge/Lasso/NNLS/Fully-invested NNLS (medie "
             "≈ 0.78). Anche la correlazione media è la più bassa "
             "(0.855) e il suo CI di IR è il più ampio. Coerentemente "
             "con la conclusione della Sezione 2: la regolarizzazione "
             "ha comprato sicurezza al prezzo di Beta. Per il "
             "production deploy, EN tuned non è la scelta giusta."),
            ("Fully-invested NNLS chiude il gap di LIVELLO ma non quello "
             "di VARIANZA. ",
             "L'IR medio è −0.045, il migliore di tutti, ma il CI è "
             "[−0.47, +0.38] e include zero. Il Beta medio (0.782) è "
             "indistinguibile da NNLS standard (0.782) e dai baseline "
             "non-regolarizzati. Σw = 1 risolve il bias, non la varianza "
             "ridotta della replica."),
            ("Il muro 78.6% è il riferimento da tenere in mente. ",
             "Qualunque modello a venire (multi-moment, constrained, "
             "liquidity-aware, regime-aware) ha questo come limite "
             "superiore di Beta. Differenze di Beta sotto il punto "
             "decimale sono fisiologiche e non vanno celebrate; "
             "differenze sopra il punto decimale sono significative."),
            ("Per le sezioni a valle, il CI bootstrap fissa la \"unità "
             "di misura\" del rumore. ",
             "Ampiezza tipica del CI: TE ~ ±0.006, IR ~ ±0.5, Beta ~ "
             "±0.06. Qualunque miglioramento più piccolo di questi "
             "valori va riportato come \"non statisticamente "
             "rilevante\"."),
        ],
    )

    # 4.6 Flags
    flags = [
        ("TE e IR puntuali sono insufficienti per il ranking dei modelli: ",
         "ogni claim che dice \"il modello X è migliore di Y per "
         "TE/IR\" senza accompagnare i CI è statisticamente vuoto. "
         "Da qui in poi ogni confronto verrà inquadrato in unità di "
         "ampiezza del CI."),
        ("Il muro teorico di R² ≈ 78.6% sul Monster Index implica "
         "Beta massimo ≈ 0.886: ",
         "modelli che si avvicinano a questo limite stanno usando bene "
         "l'universo dei futures. Modelli che si fermano sotto 0.75 "
         "stanno sprecando varianza catturabile. Le Idee 1 "
         "(multi-moment), 4 (constrained) e 5 (sparse EN+turnover) "
         "andranno valutate anche su questa scala."),
        ("HFRXGL è il vero collo di bottiglia: ",
         "R² 0.275 contro futures. È metà del target, e contribuisce "
         "alla maggior parte della varianza irriducibile. Qualunque "
         "futuro tentativo di chiudere il Beta gap dovrebbe "
         "concentrarsi qui — o cambiare la composizione del Monster, "
         "o aggiungere asset class che spannano hedge fund returns."),
        ("EN tuned è statisticamente dominato dagli altri modelli sul "
         "Beta. ",
         "Pur essendo stato selezionato come \"miglior config\" dal "
         "grid search, alla luce dei CI bootstrap risulta il modello "
         "più debole. Nelle Idee a valle (8, 8b, D, D2, E, E2) che "
         "consumano elnet_monthly come \"production model\", questa "
         "scelta è da rivedere: NNLS, Fully-invested NNLS, o "
         "Liquidity-aware (che vedremo) sono opzioni statisticamente "
         "migliori."),
        ("Fully-invested NNLS resta il vincitore sul Beta gap di "
         "LIVELLO (IR ≈ 0), ",
         "non sul Beta gap di varianza (che è strutturale)."),
        ("LEGATRUU R² = 0.512 è una piccola sorpresa: ",
         "i nostri 4 bond futures (RX1, TY1, DU1, TU2) coprono solo "
         "parte dello spettro fixed income globale. Aggiungere futures "
         "su credito o EM bond migliorerebbe il Beta più di quanto "
         "ci si aspetterebbe."),
        ("Il CI bootstrap diventa la \"unità di misura\" per tutte le "
         "sezioni successive: ",
         "miglioramenti di TE entro ±0.006 (1.3 pp di ampiezza CI) "
         "vanno trattati come rumore. Lo stesso per IR (ampiezza ~1.0) "
         "e Beta (ampiezza ~0.13)."),
    ]
    add_flag_box(doc, "Flag per il prosieguo del progetto", flags)

    add_para(doc, "")
    return doc


# ---------- SECTION 5: Multi-moment matching + HMM mixture -------------------

def section_5(doc):
    add_heading(doc, "5. Modelli alternativi: Multi-moment matching loss e "
                     "mixture HMM regime-aware", level=1)
    add_para(
        doc,
        "Dopo i modelli lineari e i vincoli di portafoglio, questa sezione "
        "esplora due direzioni più ambiziose. Idea 1 modifica la funzione "
        "di loss per inseguire non solo la varianza ma anche skewness e "
        "kurtosis del tracking difference. Idea 2 abbandona il modello "
        "singolo a favore di una miscela di replicatori, dove la scelta "
        "del modello attivo è guidata da un Hidden Markov Model addestrato "
        "su feature di rischio causali.",
    )

    # 5.1 Multi-moment
    add_heading(doc, "5.1 Idea 1 — Multi-moment matching loss", level=2)
    add_heading(doc, "Formulazione della loss", level=3)
    add_para(
        doc,
        "La TE classica minimizza solo il secondo momento del tracking "
        "difference (MSE). Il Monster Index ha però skewness ≈ −1.34 e "
        "excess kurtosis ≈ 7.58 (Sezione 1): un replicatore che "
        "azzecca la varianza ma sbaglia la forma distribuzionale produce "
        "una replica che si comporta diversamente dal target nelle code, "
        "esattamente dove conta. La loss multi-momento penalizza "
        "esplicitamente i mismatch sui terzo e quarto momento, "
        "normalizzati in modo che ciascun termine sia O(1):",
    )
    add_formula(
        doc,
        "L(w) = MSE/Var(y)  +  λ_sk · (sk_p − sk_t)² / (1 + sk_t²)  "
        "+  λ_ku · (ek_p − ek_t)² / (1 + ek_t²)  +  λ_L1 · ‖w‖₁"
    )
    add_para(
        doc,
        "Notazione: sk_t, ek_t sono skewness e excess kurtosis dei "
        "returns del target stimati sulla finestra di training; "
        "sk_p, ek_p le stesse quantità sul portafoglio Xw simulato sulla "
        "stessa finestra. La normalizzazione (1 + sk²/ek²) ai denominatori "
        "rende i termini direttamente confrontabili: λ = 1 è "
        "equal-weight fra le tre componenti.",
    )

    add_heading(doc, "Implementazione", level=3)
    add_para(
        doc,
        "Tre scelte chiave per la velocità e la stabilità:",
    )
    add_bullets(
        doc,
        [
            ("Momenti via formule NumPy: ",
             "skewness e excess kurtosis calcolati con (d³).mean()/σ³ e "
             "(d⁴).mean()/σ⁴ − 3 direttamente sull'array, evitando "
             "l'overhead di scipy.stats invocato nel loop dell'optimizer "
             "(10-50× più veloce)."),
            ("L-BFGS-B unconstrained con warm-start Ridge: ",
             "i pesi normalizzati partono dalla soluzione Ridge sulla "
             "stessa finestra, l'optimizer ha bisogno di pochissime "
             "iterazioni. Il gross cap viene applicato con una "
             "proiezione esplicita L1 dopo l'ottimizzazione (esatta per "
             "L1, costo O(k))."),
            ("Rebalance trimestrale (every 12 settimane): ",
             "il modello è ~50× più costoso del Ridge per fit; "
             "quarterly tiene il run sotto i 3 minuti per configurazione."),
        ],
    )
    add_para(
        doc,
        "Tre configurazioni di sensitivity vengono testate, identiche "
        "su tutto tranne i pesi dei momenti: λ = (0, 0) — pura TE, "
        "baseline interno; λ = (1, 1) — equal-weight; λ = (3, 3) — "
        "moment-heavy. Tutte usano il backtest engine generico con "
        "rolling window 208 settimane e normalise = False (il modello "
        "gestisce internamente lo StandardScaler).",
    )

    add_heading(doc, "Risultati", level=3)
    add_para(
        doc,
        "Tabella dei tre fit sull'OOS comune (496 settimane):",
    )
    add_table(
        doc,
        header=["Config", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett."],
        rows=[
            ["MM λ = 0 (pura TE)",  "0.0282", "−0.244", "0.852",
             "1.53", "0.028"],
            ["MM λ = 1 (balanced)", "0.0311", "−0.270", "0.817",
             "2.18", "0.052"],
            ["MM λ = 3 (heavy)",    "0.0319", "−0.300", "0.807",
             "2.40", "0.058"],
        ],
    )
    add_para(
        doc,
        "Il diagnostico del moment matching — varianza, skewness, "
        "kurtosis annualizzati della replica vs target:",
    )
    add_table(
        doc,
        header=["Modello", "Ann. Vol", "Skewness", "Excess Kurtosis"],
        rows=[
            ["Target",              "0.063", "−1.342", "7.581"],
            ["EN tuned (rif.)",     "0.044", "−1.233", "8.096"],
            ["Kalman (rif.)",       "0.044", "−1.305", "8.069"],
            ["MM λ = 0",            "0.048", "−1.456", "9.663"],
            ["MM λ = 1",            "0.047", "−1.552", "9.896"],
            ["MM λ = 3",            "0.047", "−1.550", "9.840"],
        ],
    )
    fig = ROOT / "_report_figs" / "sec5_mm_1.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Returns cumulati e distribuzioni dei tre fit "
                    "multi-momento vs target. Le tre curve cumulative "
                    "tracciano bene la direzione media, ma le distribuzioni "
                    "mostrano code più estreme della replica rispetto al "
                    "target.",
            width_cm=15,
        )

    add_heading(doc, "Lettura dei risultati", level=3)
    add_para(
        doc,
        "Il punto cruciale viene dalla seconda tabella: ",
        bold=False,
    )
    add_bullets(
        doc,
        [
            ("Il moment matching peggiora con λ crescente, non migliora. ",
             "Skewness del replicatore: EN tuned −1.23, MM λ = 0 −1.46, "
             "MM λ = 1 −1.55, MM λ = 3 −1.55. Il target è a −1.34. EN "
             "tuned è quello più vicino al target; i tre MM sono tutti "
             "*più* negativi del target, e la distanza cresce — non si "
             "riduce — passando da λ = 0 a λ = 1. Sull'eccesso di "
             "kurtosis lo stesso schema: target 7.58, EN tuned 8.10, "
             "MM 9.66–9.90."),
            ("L'optimizer overshoot: ",
             "aggiungere il termine di penalty non spinge la replica "
             "verso i momenti del target, la spinge oltre. La ragione è "
             "che skewness e kurtosis sono funzioni altamente non "
             "lineari di w; con stime di sk e ek calcolate su 208 "
             "settimane (varianza alta, soprattutto per la kurtosis), "
             "l'optimizer insegue rumore. Aumentare λ amplifica il "
             "rumore invece di azzerare il bias."),
            ("Anche TE, IR e gross exposure peggiorano: ",
             "TE sale da 0.0282 (λ = 0) a 0.0319 (λ = 3), IR da −0.244 "
             "a −0.300, gross exposure da 1.53 a 2.40. Più si chiede "
             "all'optimizer di matchare momenti, più questo leverizza il "
             "portafoglio (cercando combinazioni più estreme di "
             "long-short) e più i costi di rebalance crescono "
             "(turnover settimanale da 0.028 a 0.058)."),
            ("Anche λ = 0 (puro TE) è peggio dei modelli lineari "
             "regolarizzati: ",
             "MM λ = 0 dovrebbe equivalere a una Ridge warm-startata con "
             "proiezione L1, ma TE 0.0282 è comunque marginalmente sopra "
             "Kalman/Ridge/Lasso (≈ 0.027). Probabilmente per via della "
             "proiezione L1 con gross_cap = 2.0 che impone una "
             "configurazione meno regolarizzata di quella scelta dal "
             "grid search."),
        ],
    )
    add_para(
        doc,
        "Verdetto sul multi-moment matching: come implementato, l'idea "
        "non paga. Le tre metriche di replica peggiorano in modo "
        "monotono con λ, e la metrica che l'idea vorrebbe migliorare "
        "(il match dei momenti alti) anche essa peggiora. Il problema "
        "non è concettuale — un replicatore distribution-aware è una "
        "buona idea in linea di principio — ma di scala dei campioni: "
        "208 settimane non bastano per stimare la kurtosis del residuo "
        "in modo abbastanza stabile da darla come target a un "
        "optimizer. Estensioni naturali per future iterazioni: "
        "stimare i momenti del target con prior Bayesiani, oppure "
        "usare finestre molto più lunghe (full sample) per i momenti e "
        "rolling per i pesi.",
    )

    # 5.2 HMM
    add_heading(doc, "5.2 Idea 2 — Mixture di replicatori regime-aware (HMM)",
                level=2)
    add_para(
        doc,
        "Riprende la flag della Sezione 1.6 (correlazioni che variano fra "
        "regime calm e stress): un singolo modello lineare è obbligato "
        "a fare un compromesso tra regimi e finisce per essere "
        "sistematicamente sbagliato in entrambi. L'idea è permettere "
        "al replicatore di cambiare faccia, ma con una regola "
        "esplicita e auto-osservata: un Hidden Markov Model.",
    )
    add_para(
        doc,
        "Importante per il report: l'esecuzione di Idea 2 nell'ambiente "
        "corrente non è andata a buon fine. Ciò che segue descrive il "
        "design, la formulazione e il protocollo di valutazione; i "
        "risultati numerici non sono inclusi e dovranno essere "
        "ripresentati in una versione futura del report quando il "
        "modello sarà ri-eseguibile end-to-end.",
    )

    add_heading(doc, "Design — feature, modello, integrazione", level=3)
    add_para(
        doc,
        "Quattro feature di rischio causali (calcolate solo su passato "
        "all'interno della finestra di training) servono da input al "
        "Hidden Markov Model. Sono ispirate a indicatori macro classici "
        "(Hamilton 1989, Ang & Bekaert 2002) e — soprattutto — sono "
        "costruite solo dai futures, mai dal target:",
    )
    add_bullets(
        doc,
        [
            ("Equity vol: ",
             "deviazione standard rolling 8 settimane dei returns di "
             "ES1. Discrimina i periodi tranquilli da quelli di stress "
             "azionario."),
            ("Bond vol: ",
             "stesso indicatore su TY1. Cattura gli stress sui tassi."),
            ("Risk-on/off: ",
             "media mobile 4 settimane di (ES1 − RX1). Differenza fra "
             "equity DM e bund: una variabile flight-to-quality."),
            ("Commodity spread: ",
             "media mobile 4 settimane di (CO1 − GC1). Crude minus gold: "
             "growth vs safety."),
        ],
    )
    add_para(
        doc,
        "Sul vettore standardizzato di queste quattro feature viene "
        "addestrato un GaussianHMM (libreria hmmlearn) con covarianza "
        "diagonale (per stabilità numerica) e un floor min_covar = 10⁻³ "
        "per evitare degenerazione. Vengono testate due configurazioni: "
        "2 regimi (calm vs stress) e 3 regimi (calm / trending / "
        "crash).",
    )
    add_para(
        doc,
        "Per ciascun regime identificato dall'HMM si addestra una Ridge "
        "(α = 1) sul sottoinsieme delle settimane di training classificate "
        "in quel regime. Il modello complessivo è una miscela: a tempo t "
        "i pesi attivi sono quelli del Ridge corrispondente al regime "
        "corrente. Lo StandardScaler è gestito internamente dalla "
        "classe; il backtest engine viene invocato con normalise = False.",
    )

    add_heading(doc, "Protocollo di valutazione — causalità garantita",
                level=3)
    add_para(
        doc,
        "Lo schema di valutazione è lo stesso walk-forward dei baseline "
        "(Protocollo A della Sezione 2.5), con due cautele aggiuntive "
        "sulla causalità:",
    )
    add_bullets(
        doc,
        [
            ("Le feature di regime sono costruite usando solo dati nella "
             "finestra di training [t − K, t). ",
             "Niente leak dal test."),
            ("Il fit dell'HMM e dei Ridge per-regime usano solo la "
             "finestra di training. ",
             "Il regime per la settimana t è determinato classificando "
             "l'ultima osservazione della finestra di training (cioè "
             "t − 1), poi si predice t. Non si guarda mai t per "
             "decidere il regime di t."),
            ("Il VaR scaling è identico agli altri modelli (Cornish-Fisher "
             "su 156 settimane di replica returns), così come i costi di "
             "transazione (2 bps).",),
        ],
    )
    add_para(
        doc,
        "Una scelta di design importante: l'HMM è fittato sulle feature "
        "di rischio, non sui returns del target. Questa è una correzione "
        "rispetto a un anti-pattern frequente in letteratura, dove si "
        "addestra l'HMM direttamente sulla serie da replicare — il che "
        "introduce data leakage perché i regimi catturano esattamente la "
        "struttura che il modello dovrebbe scoprire. Qui i regimi sono "
        "scoperti in modo agnostico rispetto al target.",
    )

    add_heading(doc, "Risultati attesi vs riportabili", level=3)
    add_para(
        doc,
        "Ribadiamo: l'esecuzione del modello nell'ambiente corrente non "
        "è disponibile, quindi non riportiamo numeri. Tre fenomeni "
        "attesi dal design, da verificare quando il modello sarà "
        "ri-eseguito:",
    )
    add_bullets(
        doc,
        [
            ("La serie del posterior P(stress) dovrebbe colpire i "
             "drawdown noti del target: ",
             "2008-2009 GFC, 2011 EU debt, 2015-16 China, 2018 Q4, "
             "marzo 2020. Validazione visiva del regime detector."),
            ("Per-regime weights distinte: ",
             "in calm ci si aspetta più peso su equity (ES1, NQ1, VG1), "
             "in stress più peso su flight-to-quality (TY1, RX1, GC1) e "
             "magari pesi negativi sull'equity."),
            ("Performance complessiva: ",
             "il guadagno atteso non è grande in IR — con solo 11 "
             "futures e ~500 settimane OOS, ciascun Ridge per-regime è "
             "stimato su 100-500 osservazioni — ma l'interpretabilità è "
             "il valore aggiunto principale, oltre a una potenziale "
             "riduzione del drawdown in stress."),
        ],
    )

    box = open_methodology_box(
        doc,
        "Considerazione metodologica — perché 2 regimi sono "
        "probabilmente sufficienti"
    )
    box_para(
        box,
        "Con 11 regressori e 208 settimane di training, ciascuno dei "
        "Ridge per-regime sarebbe stimato su una media di 104 (caso 2 "
        "regimi) o 69 (caso 3 regimi) osservazioni. A 3 regimi siamo "
        "vicini al limite n / p ≈ 6, dove la varianza dei coefficienti "
        "domina e i Ridge per-regime perdono robustezza. La "
        "configurazione a 2 regimi mantiene n / p ≈ 9, ancora "
        "ragionevole. Quando il modello sarà ri-eseguito, è quindi "
        "ragionevole aspettarsi che 2 regimi battano 3 — non perché 3 "
        "siano \"sbagliati\", ma perché la nostra finestra di training "
        "non permette di stimarli bene.",
    )

    # Flags
    flags = [
        ("Multi-moment matching come implementato non paga: ",
         "tutte le metriche peggiorano con λ crescente, *inclusi* skew e "
         "kurtosis che l'idea voleva migliorare. La causa è la varianza "
         "delle stime di sk/ek su 208 settimane. Riconsiderare con "
         "finestre più lunghe per le stime di sk/ek (es. full sample) o "
         "con prior bayesiani prima di provare di nuovo."),
        ("MM crea instabilità di leva: ",
         "gross exposure passa da 1.53 (λ = 0) a 2.40 (λ = 3) e turnover "
         "raddoppia. Il VaR scaling lo limita a posteriori, ma il "
         "modello \"vuole\" leverarsi per inseguire i momenti. Questo "
         "è un anti-pattern in produzione."),
        ("MM è dominato dai modelli lineari semplici: ",
         "anche MM λ = 0 (pura TE, internamente Ridge warm-started) ha "
         "TE 0.0282 > Kalman / NNLS / EN default (0.027–0.028). Senza "
         "il vantaggio del moment matching che non si materializza, "
         "non c'è ragione di preferirlo alle alternative semplici."),
        ("HMM: il design è solido (feature causali, no leakage, "
         "regimi indipendenti dal target), ",
         "ma manca la validazione numerica. Nel run finale del progetto, "
         "ricontrollare che (i) il posterior P(stress) si allinea ai "
         "drawdown storici e (ii) le composizioni per-regime hanno "
         "senso economico (rotazione equity → bond/oro in stress)."),
        ("Per HMM: 2 regimi probabilmente meglio di 3 ",
         "per ragioni di varianza dei Ridge per-regime — n / p ≈ 9 "
         "vs 6. Quando il modello sarà ri-eseguito, partire da 2."),
        ("Idee \"sofisticate\" non sempre superano gli \"stupidi\" "
         "modelli lineari: ",
         "il Beta gap strutturale identificato nella Sezione 4 (R² ≤ "
         "78.6%) limita anche le idee non lineari. La complessità "
         "aggiuntiva paga solo se cattura varianza fuori dal linear "
         "span, e MM non sembra farlo."),
    ]
    add_flag_box(doc, "Flag per il prosieguo del progetto", flags)

    add_para(doc, "")
    return doc


# ---------- SECTION 6: Idea 4 (constrained) + Idea 5 (Elastic Net + TO) ------

def section_6(doc):
    add_heading(doc, "6. Idee a vincoli e a penalty: Idea 4 (Constrained) e "
                     "Idea 5 (Elastic Net + turnover)", level=1)
    add_para(
        doc,
        "Le due idee di questa sezione affrontano lo stesso obiettivo — "
        "rendere il replicatore operativamente compatibile con un fondo "
        "reale — con due filosofie opposte. Idea 4 codifica la struttura "
        "macro come vincoli hard nel problema di ottimizzazione: equity "
        "ed bond sleeves devono stare in intervalli prescritti, nessun "
        "asset può dominare, niente leva sostanziale. Idea 5 fa "
        "l'opposto: non vieta nulla, ma rende ogni decisione costosa "
        "tramite penalty (L1 per sparsità, L2 per stabilità, turnover "
        "per costo di esecuzione), e lascia che siano i dati a "
        "scegliere.",
    )

    # 6.1 Idea 4
    add_heading(doc, "6.1 Idea 4 — Constrained macro-allocation", level=2)

    add_heading(doc, "Motivazione e vincoli", level=3)
    add_para(
        doc,
        "Una regressione non vincolata sui returns settimanali è "
        "matematicamente ottimale ma operativamente fragile: l'optimizer "
        "concentra spesso l'esposizione su uno o due futures più "
        "correlati col target in quella specifica finestra, oppure va "
        "net-short bond in un rally obbligazionario solo perché i "
        "residui escono marginalmente più piccoli. Entrambe le scelte "
        "sono statisticamente razionali ma economicamente "
        "indefendibili. Idea 4 impone quattro vincoli simultanei, "
        "ognuno lasco abbastanza da non essere tipicamente vincolante, "
        "ma collettivamente sufficienti a escludere gli angoli "
        "patologici dello spazio dei parametri:",
    )
    add_bullets(
        doc,
        [
            ("Sleeve equity: ",
             "Σⱼ∈Equity wⱼ ∈ [0.10, 0.50]. Mai zero, mai dominante."),
            ("Sleeve bond: ",
             "Σⱼ∈Bond wⱼ ∈ [0.20, 0.50]. Ancora del portafoglio."),
            ("Leverage gross: ",
             "Σⱼ |wⱼ| ≤ 1.20. Piccolo margine sopra 1 per coperture, "
             "nessuna leva aggressiva."),
            ("Cap per strumento: ",
             "|wⱼ| ≤ 0.30. Nessuna singola posizione domina il book."),
        ],
    )
    add_para(
        doc,
        "Insieme ai vincoli c'è una piccola regolarizzazione L2 "
        "(λ = 10⁻³) per la condizionatura numerica nelle finestre di "
        "near-collinearity fra futures correlati.",
    )

    add_heading(doc, "Formulazione e protocollo", level=3)
    add_formula(
        doc,
        "min_w  ‖y_c − Xw‖² + λ‖w‖²    s.t.   vincoli sopra"
    )
    add_para(
        doc,
        "dove y_c = y − ȳ centra il target per assorbire l'intercetta. "
        "L'intercetta viene poi ricostruita esplicitamente come "
        "α = ȳ − X̄ · w per evitare il double-counting del drift. "
        "Il solver è SLSQP con warm-start feasible (punto medio dei "
        "range di sleeve distribuito sugli asset della classe), check "
        "esplicito su res.success per evitare di accettare soluzioni "
        "non-converged. Il modello entra nel backtest engine "
        "unificato con rolling window = 208 settimane, rebalance "
        "mensile (every 4 weeks), normalise = False (la classe non "
        "richiede StandardScaler perché lavora su L1/L2 sui pesi "
        "direttamente). Stesso protocollo walk-forward dei baseline, "
        "stesso OOS comune di 496 settimane.",
    )

    add_heading(doc, "Risultati", level=3)
    add_para(
        doc,
        "Sanity check sulla soluzione media (post VaR scaling):",
    )
    add_bullets(
        doc,
        [
            ("Equity sleeve medio: +0.309, ",
             "all'interno del range [0.10, 0.50], lontano da entrambi "
             "i bound — il vincolo non vincola."),
            ("Bond sleeve medio: +0.209, ",
             "praticamente sul lower bound (0.20). Il vincolo bond "
             "*è* attivo: senza di esso l'optimizer andrebbe sotto. "
             "Significato: lo sleeve bond è la copertura che il vincolo "
             "ha effettivamente \"comprato\"."),
            ("Gross medio: 0.582, ",
             "ben sotto il cap 1.20. Niente leva latente."),
            ("Max |w|: 0.181, ",
             "ben sotto il cap 0.30. Il limite per strumento è "
             "assicurazione, non meccanismo attivo."),
        ],
    )
    add_para(
        doc,
        "Performance sul comune OOS, confronto con i tre migliori "
        "modelli delle sezioni precedenti:",
    )
    add_table(
        doc,
        header=["Modello", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett."],
        rows=[
            ["EN tuned, monthly",     "0.0281", "−0.199", "0.853", "0.64", "0.018"],
            ["NNLS long-only",        "0.0275", "−0.169", "0.860", "0.64", "0.016"],
            ["Kalman (q = 10⁻⁷)",     "0.0276", "−0.154", "0.859", "0.48", "0.008"],
            ["Idea 4 (constrained)",  "0.0278", "−0.152", "0.856", "0.66", "0.012"],
        ],
    )
    fig = ROOT / "_report_figs" / "sec6_idea4_1.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Tre pannelli: returns cumulati netti vs target; "
                    "evoluzione delle sleeve equity / bond / commodity nel "
                    "tempo (le sleeve restano nei range prescritti per "
                    "tutto il sample); evoluzione del gross exposure (ben "
                    "sotto il cap 1.20).",
            width_cm=14,
        )
    add_para(
        doc,
        "Tre conclusioni nette:",
    )
    add_bullets(
        doc,
        [
            ("IR netto −0.152, il migliore del confronto a 4 candidati. ",
             "Marginalmente sopra Kalman (−0.154, di fatto un pareggio), "
             "ma con uno scarto piccolo rispetto all'ampiezza del CI "
             "bootstrap della Sezione 4 (~0.5 sulla IR): la differenza è "
             "rumore statistico. La conclusione robusta è \"Idea 4 e "
             "Kalman sono indistinguibili sul fronte IR\"."),
            ("Gli unici vincoli vincolanti sono i lower bound dei sleeve. ",
             "L'equity è a +0.31 (in range), gross 0.58 (cap inattivo), "
             "max-pos 0.18 (cap inattivo). Solo il bond floor (0.20) "
             "tocca il vincolo. Questo è importante: significa che i "
             "vincoli funzionano come assicurazione, non come "
             "steering. L'optimizer prenderebbe configurazioni simili "
             "anche senza, *ma* i vincoli garantiscono di non andare "
             "fuori in regimi che non abbiamo ancora visto."),
            ("Turnover 0.012, fra Kalman (0.008) e NNLS (0.016). ",
             "Idea 4 paga la rigidità con un po' più di churn rispetto "
             "alla dinamica continuativa del Kalman, ma è ancora 1.5× "
             "più stabile dell'EN tunato."),
        ],
    )
    add_para(
        doc,
        "Punto di forza specifico, non visibile nei numeri: la "
        "portafoglio è descrivibile in una frase a un PM — \"equity "
        "intorno al 30%, bond intorno al 20%, gross al 60%, nessuna "
        "posizione sopra il 20% in modulo\". È la proprietà "
        "interpretativa che nessun modello puramente statistico ha, e "
        "fa di Idea 4 un candidato di produzione plausibile.",
    )

    # 6.2 Idea 5
    add_heading(doc, "6.2 Idea 5 — Elastic Net + Turnover Penalty", level=2)

    add_heading(doc, "Formulazione", level=3)
    add_para(
        doc,
        "Tre penalty si aggiungono al fit standard, ciascuna con un "
        "ruolo economico distinto:",
    )
    add_formula(
        doc,
        "L(w) = (1/2n) · ‖y − Xw‖²  +  α · ρ · ‖w‖₁  "
        "+  ½ · α · (1−ρ) · ‖w‖²  +  λ_TO · ‖w − w_{t−1}‖₁"
    )
    add_bullets(
        doc,
        [
            ("α · ρ · ‖w‖₁ (L1, sparsità): ",
             "geometricamente la palla L1 ha angoli sugli assi → la "
             "soluzione tende a esibire weight esattamente zero su "
             "alcuni regressori. Economicamente: ogni future incluso "
             "deve \"giustificarsi\"."),
            ("½α(1−ρ)‖w‖² (L2, stabilità): ",
             "scoraggia coefficienti grandi in qualsiasi direzione, "
             "smussa la loss surface nelle zone di multicollinearità "
             "(critico per i 5 equity futures, ≥ 90% correlati fra "
             "loro)."),
            ("λ_TO · ‖w − w_{t−1}‖₁ (turnover): ",
             "penalty applicata nello spazio reale dei pesi (non "
             "normalizzato), espressa in unità di costo di portafoglio. "
             "Calibrata a ~10% della contribuzione tipica della loss "
             "MSE — turnover scoraggiato ma non vietato."),
        ],
    )

    add_heading(doc, "Tuning automatico e protocollo", level=3)
    add_para(
        doc,
        "I due parametri (α, ρ) vengono selezionati da ElasticNetCV "
        "sulle prime 2·window_size = 156 settimane di calibrazione, con "
        "TimeSeriesSplit(n_splits = 5) e griglie alphas = "
        "np.logspace(−6, −3, 30), L1_RATIOS = [0.1, 0.3, 0.5, 0.7, 0.9]. "
        "Il backtest poi rotola su un rolling window di 78 settimane "
        "(≈ 1.5 anni); ogni 26 settimane (~6 mesi) si rifa la "
        "calibrazione CV sulle ultime 156 settimane.",
    )
    add_para(
        doc,
        "Tre scelte di disegno meritano nota esplicita:",
    )
    add_bullets(
        doc,
        [
            ("ALPHA_MAX = 10⁻⁴ come tetto duro per α: ",
             "ElasticNetCV può scegliere α fino a 10⁻², ma sopra ~10⁻⁴ "
             "il termine L1 domina la (1/2n)‖y−Xw‖² (var del Monster "
             "settimanale ≈ 10⁻⁴) e la soluzione collassa al vettore "
             "nullo. Da quel punto, il warm-start a zero combinato con "
             "la turnover penalty inchioda l'intero loop a zero. Il "
             "cap è il floor di stabilità."),
            ("Warm-start direttamente da enet_cv.coef_, non da un "
             "refit su finestra ridotta: ",
             "rifittare ElasticNet su 78 settimane con lo stesso α "
             "produce una soluzione *più densa* (meno dati = meno "
             "regolarizzazione effettiva) e seederebbe il rolling loop "
             "con tutti gli 11 futures attivi, distruggendo la "
             "sparsità che la CV aveva identificato."),
            ("Rebalance settimanale, con la turnover penalty come "
             "freno: ",
             "a differenza di Idea 4 (mensile per scelta), Idea 5 "
             "rebalancia ogni settimana ma quasi nessuna porta a "
             "movimenti effettivi. È la differenza fra \"rebalance "
             "schedulato\" e \"rebalance condizionale\"."),
        ],
    )

    box = open_methodology_box(
        doc,
        "Nota di confronto con i protocolli del notebook"
    )
    box_para(
        box,
        "Idea 5 NON utilizza il backtest engine unificato (cella 16). "
        "Le ragioni sono pragmatiche (la turnover penalty richiede di "
        "conoscere w_{t−1} dentro la loss function, integrazione "
        "non triviale) ma producono alcune differenze rispetto al "
        "resto del notebook che è importante esplicitare per leggere "
        "correttamente i risultati:",
    )
    box_bullets(
        box,
        [
            ("Window = 78 settimane (1.5 anni), ",
             "contro 208 (4 anni) di tutti gli altri modelli post-tuning. "
             "Il sample di training è dimezzato."),
            ("Costi di transazione TC_BPS = 5 bps, ",
             "contro 2 bps del resto del notebook. Idea 5 sconta un "
             "costo di esecuzione più realistico."),
            ("Loop di backtest custom, non walk-forward del engine "
             "comune: ",
             "il rolling è gestito internamente alla cella; le metriche "
             "(R² OOS, hit rate, TE annualizzata) sono calcolate "
             "direttamente sul vettore predicted_returns vs "
             "actual_returns senza passare da report()."),
            ("La conseguenza è che IR e TE di Idea 5 non sono "
             "direttamente confrontabili con quelle delle sezioni "
             "precedenti. ",
             "Su un protocollo a parità (window 208, TC 2bps, engine "
             "unificato) i numeri potrebbero spostarsi. Vanno letti "
             "come performance di un modello a sé, non come "
             "miglioramento numerico di Idea 4."),
        ],
    )

    add_heading(doc, "Risultati", level=3)
    add_para(
        doc,
        "Output di calibrazione: α* = 10⁻⁴ (capped), ρ* = 0.30, "
        "warm-start con 10 features non-zero su 11. λ_TO = 9.03 · 10⁻⁶ "
        "(0.10 × loss tipica per osservazione 9.03 · 10⁻⁵).",
    )
    add_para(
        doc,
        "Performance:",
    )
    add_table(
        doc,
        header=["Metrica", "Valore"],
        rows=[
            ["Correlation",                "0.859"],
            ["TE annualizzata",            "2.92%"],
            ["R² OOS",                     "0.736"],
            ["Hit rate direzionale",       "84.5%"],
            ["IR (Idea 5 protocollo)",     "+0.028"],
            ["Gross exposure",             "0.758"],
            ["Turnover settimanale",       "0.03%"],
            ["Costo annuo TC stimato",     "0.001%"],
            ["Active futures (medio)",     "8.6 / 11"],
            ["52w rolling corr — min",     "0.599"],
            ["52w rolling corr — % > 0.5", "100%"],
        ],
    )
    fig = ROOT / "_report_figs" / "sec6_idea5_1.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Pannelli diagnostici di Idea 5: returns cumulati, "
                    "evoluzione del gross exposure (stabile attorno a "
                    "0.76), rolling correlation 52w (mai sotto 0.6), "
                    "composizione asset class e pesi medi per future.",
            width_cm=14,
        )
    add_para(
        doc,
        "Pesi medi per future (sintesi grafica dal diagnostic):",
    )
    add_table(
        doc,
        header=["Future", "Peso medio", "% settimane attive",
                "Asset class"],
        rows=[
            ["TY1",  "+0.241", "100%", "Bond"],
            ["ES1",  "+0.147", "100%", "Equity DM"],
            ["DU1",  "−0.089", "49%",  "Bond"],
            ["VG1",  "+0.070", "100%", "Equity DM"],
            ["TP1",  "+0.068", "100%", "Equity DM"],
            ["RX1",  "+0.060", "100%", "Bond"],
            ["CO1",  "+0.032", "100%", "Commodity"],
            ["GC1",  "+0.031", "100%", "Commodity"],
            ["NQ1",  "+0.019", "100%", "Equity DM"],
            ["LLL1", "+0.000", "10.5%", "Equity EM"],
            ["TU2",  "+0.000", "0%",   "Bond"],
        ],
    )

    add_heading(doc, "Lettura", level=3)
    add_para(
        doc,
        "Quattro letture esplicite, in ordine di importanza:",
    )
    add_bullets(
        doc,
        [
            ("Sparsità genuina, non vincolata: ",
             "TU2 è permanentemente spento (0% delle settimane), LLL1 "
             "quasi sempre (10.5%). Il modello \"capisce\" che la "
             "duration corta US è già catturata da TY1 e RX1, e che "
             "l'EM equity è dominata dall'equity DM. Sparsità "
             "data-driven prodotta dalla geometria L1 sulla matrice di "
             "covarianza dei futures."),
            ("Turnover ~40× inferiore a Idea 4: ",
             "0.03% settimanale vs 1.2%. La turnover penalty fa "
             "esattamente quello che promette. Il chart del gross "
             "exposure (vedi figura) è quasi una retta — il modello "
             "rebalancia formalmente ogni settimana ma quasi nessuna "
             "rebalance produce movimenti significativi."),
            ("DU1 short ~49% del tempo: ",
             "il modello tiene DU1 (bond tedesco a breve) short "
             "metà del tempo come hedge contro la concentrazione di "
             "duration in TY1 + RX1 quando la curva si appiattisce. "
             "Nessuno l'ha imposto: è una covertura emergente dalla "
             "regolarizzazione."),
            ("Gap di livello strutturale ~5–7%: ",
             "il cumulato della replica gira ~5-7% sotto il target nel "
             "12y. È il prezzo della sparsità — gross 0.76 < 1 — e non "
             "è un bug. È coerente con il muro R² ≤ 78.6% identificato "
             "nella Sezione 4."),
        ],
    )

    # 6.3 Idea 4 vs Idea 5
    add_heading(doc, "6.3 Idea 4 vs Idea 5 — due filosofie a confronto",
                level=2)
    add_para(
        doc,
        "Le due idee replicano lo stesso target con la stessa qualità "
        "statistica di base (correlazione ≈ 0.86, TE ≈ 3%) ma con "
        "filosofie matematiche opposte:",
    )
    add_table(
        doc,
        header=["Asse", "Idea 4 (Constrained SLSQP)",
                "Idea 5 (Elastic Net + TO)"],
        rows=[
            ["Filosofia",
             "Vincoli hard, regolarizzazione esogena",
             "Penalty soft, regolarizzazione endogena"],
            ["Cosa è proibito",
             "Configurazioni fuori dai range",
             "Niente: tutto è penalizzato in base al costo"],
            ["Solver",
             "SLSQP (QP con linear constraints)",
             "L-BFGS-B su loss composta"],
            ["Window training",
             "208 settimane",
             "78 settimane (+ 156 di calibrazione CV)"],
            ["Hyperparameters",
             "Zero (window e bound fissi)",
             "α, ρ scelti dalla CV; ALPHA_MAX e λ_TO calibrati a mano"],
            ["Rebalance",
             "Mensile (every 4 weeks)",
             "Settimanale, ma penalizzato"],
            ["Sparsità",
             "Nessuna (11/11 attivi)",
             "Sì (8.6/11 attivi)"],
            ["Turnover settimanale",
             "0.012",
             "0.0003"],
            ["Robustezza",
             "Per costruzione (vincoli)",
             "Per selezione (penalty)"],
        ],
    )
    add_para(
        doc,
        "Il punto centrale del confronto: nessuna delle due strategie "
        "domina l'altra. Sono ottimali per scenari di failure mode "
        "diversi.",
    )
    add_bullets(
        doc,
        [
            ("Idea 4 è robusta contro regimi nuovi e contro errori "
             "sistematici del dato: ",
             "non importa cosa dicono i residui, i vincoli garantiscono "
             "una composizione macro sensata. Se i futures hanno un "
             "outlier di crisi nel sample di training (ad esempio una "
             "settimana di stress sub-rappresentata), l'optimizer non "
             "potrà fuggire dai bound. Costo: 40× più turnover di Idea "
             "5, qualche basis point di TE in più."),
            ("Idea 5 è robusta contro la sovra-allocazione e i costi: ",
             "la penalty L1 garantisce sparsità, la turnover penalty "
             "garantisce stabilità in tempo, la CV regola l'aggressività "
             "complessiva ai dati. Costo: la sparsità chiede un gross "
             "exposure inferiore al ideale, da cui il level gap di "
             "5–7% sul cumulato."),
            ("Confronto IR ingannevole: ",
             "Idea 5 riporta IR +0.028, Idea 4 IR −0.152. Differenza "
             "apparente di 0.18, enorme. Ma usano protocolli diversi "
             "(window, TC, engine) e i numeri non sono direttamente "
             "comparabili. L'unico confronto a parità è quello di "
             "correlazione (0.856 vs 0.859) e qualitativo "
             "(composizione, turnover, sparsità) — su questi i due "
             "modelli sono pari."),
        ],
    )

    # 6.4 Flags
    flags = [
        ("Idea 4 è statisticamente indistinguibile da Kalman (IR "
         "−0.152 vs −0.154): ",
         "i due modelli sono pari sul fronte tracking. Il "
         "differenziatore è la composizione: Idea 4 garantisce "
         "componenti macro non patologiche; Kalman è più "
         "cost-efficient (turnover 0.008 vs 0.012). La scelta fra i "
         "due dipende dal contesto operativo, non da una metrica "
         "decisiva."),
        ("I vincoli di Idea 4 sono \"insurance, not steering\": ",
         "solo il bond floor (0.20) è attivo nella soluzione media. "
         "Equity sleeve e gross cap sono lontani dai bound. "
         "Significa che il valore aggiunto di Idea 4 si manifesterà "
         "in regimi futuri stressati, non sul backtest 2011-2021."),
        ("Idea 5 produce sparsità genuina (TU2 sempre spento, LLL1 "
         "quasi sempre): ",
         "questa è informazione utile per il design degli universi "
         "futuri — TU2 (duration corta US) e LLL1 (EM equity) sono "
         "potenzialmente eliminabili dall'universo senza degrado "
         "della replica."),
        ("La turnover penalty di Idea 5 funziona straordinariamente "
         "bene: ",
         "0.03% settimanale vs 1.2% di Idea 4 (40×). Se il prossimo "
         "step di lavoro fosse trasferire questa tecnica al backtest "
         "engine unificato, su window = 208 e TC = 2bps, potremmo "
         "stimare di quanto la turnover penalty migliora un Kalman/"
         "Ridge in maniera direttamente comparabile."),
        ("Idea 5 non vive nel backtest engine unificato: ",
         "window 78 (non 208), TC 5bps (non 2), loop custom. I suoi "
         "numeri non sono nelle tabelle di confronto a parità del "
         "resto del notebook. Per una versione finale del progetto, "
         "questa integrazione è il passo naturale: la sparsità + "
         "turnover penalty sono ingredienti generici trasferibili a "
         "tutti gli altri modelli."),
        ("DU1 short ~49% del tempo in Idea 5: ",
         "comportamento emergente di hedge contro la concentrazione "
         "TY1+RX1. Da tenere a mente per il design di replicatori "
         "long-only (NNLS, Fully-invested NNLS): se DU1 è strutturalmente "
         "negativo nel fit ottimo, l'imposizione w ≥ 0 elimina un "
         "hedge importante. Spiega in parte perché NNLS fully-invested "
         "ha TE leggermente più alta."),
        ("Il muro R² ≤ 78.6% spiega anche il gap di livello di Idea "
         "5: ",
         "gross 0.76 e 7% di sottoperformance cumulata sono coerenti "
         "con il limite teorico identificato nella Sezione 4. Nessuna "
         "modifica al solo modello chiuderà quel gap; serve cambiare "
         "l'universo o passare a non-lineare."),
    ]
    add_flag_box(doc, "Flag per il prosieguo del progetto", flags)

    # Box: future improvement — wrap Idea 5 in unified backtest
    box = open_methodology_box(
        doc,
        "Sviluppi futuri — integrazione di Idea 5 nel backtest engine "
        "unificato"
    )
    box_para(
        box,
        "Idea 5 vive oggi in una cella che gestisce il proprio loop di "
        "backtest e i propri parametri (window 78 settimane, TC 5 bps, "
        "nessun VaR scaling). Questo limita la confrontabilità dei suoi "
        "numeri con il resto del notebook. Per una versione consolidata "
        "del progetto, l'integrazione nel backtest engine unificato è "
        "il passo naturale. Si distinguono tre livelli di intervento, "
        "in ordine di costo crescente.",
    )
    box_subhead(box, "Livello 1 — Solo riallineamento dei parametri "
                     "(costo ~mezz'ora)")
    box_bullets(
        box,
        [
            ("Cambiare window_size da 78 a 208, ",
             "stessa rolling window dei modelli post-tuning."),
            ("Allineare TC_BPS a 2e-4, ",
             "coerente col resto del notebook."),
            ("Aggiungere VaR Cornish-Fisher scaling con cap UCITS 20%, ",
             "riusando lo stesso pattern del backtest engine."),
            ("Ricalibrare λ_TO: ",
             "il valore corrente è dimensionato sulla loss tipica con "
             "window=78; con window=208 la varianza media della loss "
             "cambia e λ_TO va ricalcolato per mantenere il rapporto "
             "~10% loss/turnover-penalty."),
        ],
    )
    box_para(
        box,
        "Risultato atteso: numeri di Idea 5 nella stessa scala degli "
        "altri modelli; potremmo finalmente confrontare TE e IR "
        "direttamente con EN tunato, NNLS e Kalman. Il loop custom resta, "
        "ma i risultati sono interpretabili nello stesso quadro.",
    )
    box_subhead(box, "Livello 2 — Wrap completo come model factory "
                     "sklearn-style (costo significativo)")
    box_para(
        box,
        "Trasformare Idea 5 in una classe wrapper sklearn-compatibile "
        "che esponga fit/coef_/intercept_ e venga consumata da "
        "backtest() come qualunque altro modello. La difficoltà tecnica "
        "principale è che la turnover penalty richiede conoscere "
        "w_{t−1} dentro la loss function, quindi il wrapper deve "
        "mantenere stato attraverso le chiamate di fit successive — "
        "esattamente la stessa sfida risolta per il Kalman con "
        "backtest_adaptive_kalman.",
    )
    box_bullets(
        box,
        [
            ("Una implementazione possibile è una classe "
             "TurnoverPenalizedElasticNet con attributo "
             "self.prev_weights, ",
             "inizializzato a zero, aggiornato a self.coef_ dopo "
             "ogni fit. Il backtest engine deve garantire di "
             "riutilizzare la stessa istanza (model_factory() che "
             "ritorna il singleton, non una nuova istanza ogni "
             "rebal)."),
            ("In alternativa, ",
             "estendere backtest() con un parametro stateful_factory "
             "che indichi se mantenere stato fra chiamate, oppure "
             "passare prev_w come argomento esplicito tramite "
             "weight_post_fn."),
        ],
    )
    box_para(
        box,
        "Risultato atteso: la combinazione \"L1 + L2 + turnover\" "
        "diventa un ingrediente trasferibile a tutti gli altri "
        "modelli — non solo Idea 5 in isolamento. Potremmo "
        "applicarla a Ridge, Lasso, NNLS, persino agli output di "
        "Kalman. Questa è la cosa più interessante: la turnover "
        "penalty di Idea 5 è già documentata come straordinariamente "
        "efficace (40× meno turnover di Idea 4 con stessa "
        "correlation); trasferirla aggiunge un degree of freedom "
        "ortogonale a tutti i modelli precedenti.",
    )
    box_subhead(box, "Livello 3 — Ridisegno della loss penalty come "
                     "ingrediente universale")
    box_para(
        box,
        "Estensione naturale del Livello 2: definire una libreria di "
        "penalty (L1, L2, turnover, liquidity-aware) componibili "
        "tramite un'API uniforme. Ogni modello del notebook diventa "
        "una scelta di (base estimator) × (composizione di penalty) × "
        "(rebalance schedule). È il design implicito a cui le sezioni "
        "successive (Idea A liquidity-aware) stanno già convergendo: "
        "il prossimo passo logico è formalizzarlo.",
    )
    box_para(
        box,
        "Per il presente progetto è una direzione di lavoro futuro "
        "esplicita; per la consegna corrente, il Livello 1 è il "
        "compromesso ragionevole (numeri direttamente comparabili "
        "con relativa poca fatica).",
    )

    add_para(doc, "")
    return doc


# ---------- SECTION 7: Liquidity-aware + Adaptive band + Combined ------------

def section_7(doc):
    add_heading(doc, "7. Estensioni production-oriented: Liquidity-aware, "
                     "Adaptive band, combinato A+B", level=1)
    add_para(
        doc,
        "Le sezioni precedenti hanno costruito un replicatore credibile "
        "\"su carta\". Le due idee di questa sezione si occupano dei "
        "due ostacoli che separano un backtest da una strategia "
        "consegnabile in produzione: il fatto che i costi di esecuzione "
        "non sono uniformi (alcuni futures sono enormemente più liquidi "
        "di altri) e il fatto che un calendario di rebalance fisso "
        "spende soldi anche quando non c'è informazione nuova. Idea A "
        "interviene a livello di fit, Idea B a livello di policy di "
        "ribilanciamento. La sezione si chiude con il loro combinato e "
        "con una discussione esplicita di come integrarli con i modelli "
        "delle sezioni precedenti.",
    )

    # 7.1 Idea A
    add_heading(doc, "7.1 Idea A — Liquidity-aware execution penalty",
                level=2)

    add_heading(doc, "Motivazione e formulazione", level=3)
    add_para(
        doc,
        "Sul nostro universo di futures i costi di esecuzione non sono "
        "affatto uniformi. Treasury e S&P (TY1, ES1) si trattano per "
        "circa 1 bp; il Bund e i bond brevi (RX1, TU2) per 1.5 bps; "
        "l'oro e il Nasdaq (GC1, NQ1) per 2 bps; il Topix (TP1) per 4 "
        "bps; l'MSCI Emerging (LLL1) per 6 bps. Un fit che ignora queste "
        "differenze concentra liberamente l'esposizione sui contratti "
        "con prior bid-ask alto se questo riduce marginalmente l'errore "
        "di tracking — ma in produzione quella scelta costa fino a 6× "
        "rispetto all'alternativa liquid.",
    )
    add_para(
        doc,
        "Idea A formalizza il costo nella funzione obiettivo. Per ogni "
        "asset j definiamo due parametri:",
    )
    add_bullets(
        doc,
        [
            ("κⱼ ", "(linear bid-ask spread, in frazione): proporzionale "
             "al prior di liquidità dell'asset (1 bp ÷ 6 bps secondo il "
             "dizionario LIQUIDITY_PRIOR_BPS). Una modulazione "
             "anti-ciclica moltiplica per ~1.4 nei periodi di alta vol "
             "recente, ~0.7 in calm — il bid-ask si allarga nei "
             "regimi stressati."),
            ("λⱼ ", "(market impact, non-lineare): proporzionale a 4× κⱼ "
             "(IMPACT_PRIOR_BPS). Modella che ordini più grandi "
             "introducono una concessione di prezzo crescente con la "
             "size."),
        ],
    )
    add_para(
        doc,
        "La loss minimizzata a ogni rebalance è:",
    )
    add_formula(
        doc,
        "L(w) = MSE(w)  +  α_calib · Σⱼ [ κⱼ · |Δwⱼ|  +  λⱼ · |Δwⱼ|^1.5 ]"
    )
    add_para(
        doc,
        "dove Δw = w − w_prev (turnover rispetto ai pesi della "
        "rebalance precedente). Il primo termine penalizza la "
        "distanza in scala lineare (bid-ask), il secondo aggiunge una "
        "non-linearità (^1.5) che codifica l'impact. La penalty totale "
        "è scalata da α_calib, una costante di auto-calibrazione che "
        "dimensiona la penalty a una frazione fissa (1/100) del MSE "
        "iniziale per un movimento di riferimento Δw_ref = 0.005 per "
        "asset. Questo rende penalty_intensity (= 1.0 nei nostri run) "
        "un parametro scale-invariant e direttamente interpretabile.",
    )
    add_para(
        doc,
        "L'ottimizzazione è L-BFGS-B con warm-start Ridge (α = 10⁻⁴), "
        "gradiente analitico, valore-assoluto smussato (eps = 10⁻¹²) "
        "per la differenziabilità. Tutto su scala raw — niente "
        "StandardScaler — perché la penalty è interpretabile solo in "
        "unità di peso reale.",
    )

    add_heading(doc, "Risultati", level=3)
    add_para(
        doc,
        "Confronto con l'Elastic Net tuned, sul comune OOS (496 "
        "settimane):",
    )
    add_table(
        doc,
        header=["Modello", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett."],
        rows=[
            ["EN tuned (flat 2 bps)", "0.0281", "−0.199", "0.853", "0.64",
             "0.0182"],
            ["Liquidity-aware V6",     "0.0276", "**−0.069**", "0.859",
             "**0.77**", "**0.0045**"],
        ],
    )
    add_para(
        doc,
        "Tre risultati netti, ciascuno migliore di qualunque altro "
        "modello del notebook:",
    )
    add_bullets(
        doc,
        [
            ("IR netto −0.069: ",
             "il miglior valore puntuale fra tutti i modelli lineari "
             "del progetto, compresi Kalman (−0.154), NNLS (−0.169) e "
             "Idea 4 (−0.152). Statisticamente è ancora dentro la banda "
             "CI bootstrap della Sezione 4 (IR ∈ [−0.57, +0.35]), ma "
             "è l'unico modello il cui valore centrale è praticamente "
             "indistinguibile da zero."),
            ("Gross exposure 0.77, ",
             "il più alto fra i modelli a turnover basso. Sale di 13 "
             "punti rispetto a EN tuned (0.64), avvicinandosi al "
             "fully-invested. Questo spiega il salto di IR: la replica "
             "è meglio scalata verso il target, sotto-investe meno "
             "drift."),
            ("Turnover 0.0045/sett. (4× più basso di EN tuned, ",
             "1.7× più basso di Kalman). La penalty κⱼ disincentiva i "
             "movimenti più di quanto incentivi i pesi sui futures "
             "illiquidi. Le revisioni periodiche dei pesi si "
             "concentrano su TY1 e ES1 (i due asset più liquidi), dove "
             "il rapporto segnale-rumore è alto e il costo di "
             "spostarsi è basso."),
        ],
    )

    fig = ROOT / "_report_figs" / "sec7_liq_compshift_1.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Composition shift: pesi medi |w| per future sotto "
                    "EN tuned vs Liquidity-aware. La penalty redistribuisce "
                    "peso verso TY1/ES1 (liquidi), riducendo l'esposizione "
                    "su DU1/NQ1 (meno liquidi) e tenendo bassi LLL1/CO1.",
            width_cm=14,
        )

    add_para(
        doc,
        "Il composition shift è la firma diagnostica di Idea A. La "
        "tabella in figura ordina i futures dal più liquido (prior 1 "
        "bp) al meno liquido (prior 6 bps). La colonna Δ% mostra come "
        "cambia il peso medio passando da EN tuned a Liquidity-aware: "
        "TY1 (+94%), ES1 (+39%), RX1 (+110%) salgono; NQ1 (−35%) e DU1 "
        "(−21%) scendono. LLL1, paradossalmente, sale (+286%) ma in "
        "valore assoluto resta minuscolo (da 0.7% a 2.8%) — è un effetto "
        "di scala del calcolo percentuale su pesi quasi nulli, non un "
        "movimento operativo significativo.",
    )

    # 7.2 Idea B
    add_heading(doc, "7.2 Idea B — Adaptive no-trade band", level=2)

    add_heading(doc, "Motivazione e regola di decisione", level=3)
    add_para(
        doc,
        "Un calendario di rebalance fisso (mensile, settimanale) "
        "spende soldi anche quando non c'è una buona ragione per "
        "rebalanciare. Una soluzione naïve sarebbe \"non rebalanciare "
        "se il turnover atteso è sotto una soglia fissa\", ma una "
        "soglia fissa è cieca alle condizioni di mercato: in stress il "
        "modello dovrebbe muoversi di più, in calm di meno. Idea B "
        "esegue il rebalance se e solo se il guadagno atteso in "
        "Tracking Error supera il costo, moltiplicato per un soglia "
        "adattiva θ_t.",
    )
    add_formula(
        doc,
        "rebalance ⟺ E[ΔTE] · h  >  θ_t · (TC_BPS · turnover)"
    )
    add_para(
        doc,
        "dove h = 4 settimane (orizzonte su cui annualizzare il "
        "guadagno) e θ_t è la soglia adattiva. La forma di θ_t è:",
    )
    add_formula(
        doc,
        "θ_t = θ_base  ·  √(vol_baseline/vol_recent)  ·  "
        "√(TE_baseline/TE_recent)  ·  patience"
    )
    add_bullets(
        doc,
        [
            ("vol_factor: ",
             "alta vol attuale → soglia bassa, il modello DEVE "
             "muoversi quando il regime si rompe."),
            ("perf_factor: ",
             "TE attuale alta rispetto al baseline → soglia bassa, "
             "il modello DEVE muoversi quando il tracking peggiora."),
            ("patience: ",
             "se non si è rebalanciato da N settimane, la soglia si "
             "abbassa progressivamente; dopo max_skip = 12 settimane "
             "il rebalance è forzato. Protezione contro \"freeze\" "
             "patologici dei pesi."),
        ],
    )
    add_para(
        doc,
        "θ_base = 1.5 di default: chiediamo che il guadagno sia 1.5× "
        "il costo prima di muovere il portafoglio. Il valore è "
        "deliberatamente più alto di 1 per coprire l'incertezza nella "
        "stima del guadagno.",
    )

    add_heading(doc, "Integrazione con il backtest engine — più modelli "
                     "sottostanti", level=3)
    add_para(
        doc,
        "La funzione backtest_adaptive() è generica nel parametro "
        "model_factory, quindi la band si compone naturalmente sopra "
        "qualsiasi modello sklearn-style. La applichiamo a tre "
        "candidati:",
    )
    add_bullets(
        doc,
        [
            ("EN tuned: ",
             "factory standard, scelta di riferimento del notebook."),
            ("NNLS long-only: ",
             "factory già esistente nella cella 26, integrata "
             "direttamente — normalise=True (default) preserva la "
             "non-negatività via unscaling beta_raw = coef · σ_y/σ_x "
             "(σ > 0)."),
            ("Kalman filter (stateful): ",
             "il Kalman non si esprime come factory perché è uno stato "
             "che evolve, non un modello che si ri-fitta da capo ogni "
             "rebalance. Per integrarlo serve una funzione dedicata "
             "(backtest_adaptive_kalman) che mantiene (β, P) attraverso "
             "tutte le settimane e applica la band solo al passo di "
             "deploy: ogni 4 settimane si decide se il β filtrato "
             "corrente sostituisce il last_w o se quest'ultimo resta. "
             "Il filtro si aggiorna comunque ogni settimana sulla yt "
             "realizzata."),
        ],
    )

    fig = ROOT / "_report_figs" / "sec7_band_diagnostic_1.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Diagnostica della band su Elastic Net tuned. "
                    "Pannello superiore: θ_t nel tempo, con i puntini "
                    "rossi che marcano i rebalance effettivamente "
                    "eseguiti. Pannello centrale: guadagno atteso vs "
                    "hurdle θ_t·cost. Pannello inferiore: volatilità "
                    "rolling 13w del target (driver di vol_factor).",
            width_cm=14,
        )

    add_heading(doc, "Risultati per modello sottostante", level=3)
    add_table(
        doc,
        header=["Modello", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett.", "% rebal eseguiti"],
        rows=[
            ["EN tuned (sempre)",       "0.0281", "−0.199", "0.853",
             "0.64", "0.0182", "100%"],
            ["EN + adaptive band",      "0.0280", "−0.203", "0.854",
             "0.63", "0.0137", "55.6%"],
            ["Kalman (sempre)",         "0.0276", "−0.154", "0.859",
             "0.48", "0.0077", "100%"],
            ["Kalman + adaptive band",  "0.0279", "−0.146", "0.855",
             "0.48", "0.0041", "70.2%"],
            ["NNLS (sempre)",           "0.0275", "−0.169", "0.860",
             "0.64", "0.0160", "100%"],
            ["NNLS + adaptive band",    "0.0275", "−0.161", "0.860",
             "0.63", "0.0121", "62.1%"],
        ],
    )
    add_para(
        doc,
        "Conclusioni nette sui tre confronti:",
    )
    add_bullets(
        doc,
        [
            ("Il turnover scende sempre, ",
             "del 25% su EN (0.0182 → 0.0137), del 47% su Kalman "
             "(0.0077 → 0.0041), del 24% su NNLS (0.0160 → 0.0121). "
             "La band fa il suo lavoro: salta i rebalance che non "
             "valgono il costo."),
            ("Il TE non cambia statisticamente: ",
             "in nessun caso la differenza eccede 0.5 punti base, ben "
             "dentro l'ampiezza del CI bootstrap (~0.6 pp). La "
             "tracking quality è invariante."),
            ("L'IR è confuso — leggermente peggio su EN, meglio su "
             "Kalman e NNLS, ",
             "ma sempre dentro ±0.04, di nuovo dentro l'incertezza "
             "bootstrap (~0.5 di ampiezza CI per l'IR). "
             "Statisticamente, la band non sposta l'IR."),
            ("Verdetto: la band riduce i costi senza degradare la "
             "tracking quality, ma il guadagno economico è piccolo "
             "(2-9 bps/anno a 2 bps/turnover). ",
             "Non è un vantaggio robusto come quello che la Liquidity-"
             "aware estrae al fronte di esecuzione."),
        ],
    )

    # 7.3 Combined A+B
    add_heading(doc, "7.3 Combinato Liquidity-aware + Adaptive band",
                level=2)
    add_para(
        doc,
        "La composizione naturale: prima la liquidity-aware fitter "
        "produce pesi economicamente più sostenibili (κⱼ-aware); poi "
        "la band decide se commitarli o tenere quelli vecchi. "
        "L'implementazione è una versione di backtest_liquidity_adaptive "
        "che fonde i due meccanismi e usa il per-asset cost al posto "
        "di TC_BPS · turnover anche nella regola di decisione.",
    )
    add_para(
        doc,
        "Risultati sul comune OOS:",
    )
    add_table(
        doc,
        header=["Modello", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett."],
        rows=[
            ["Liquidity-aware (solo A)",       "0.0276", "−0.069", "0.859",
             "0.77", "0.0045"],
            ["Liq + Adaptive (A+B, EN base)",  "0.0278", "−0.154", "0.857",
             "0.77", "0.0046"],
        ],
    )
    add_para(
        doc,
        "Il combinato è essenzialmente identico al solo A in TE, "
        "correlation, gross exposure e turnover. **Ma l'IR netto scende "
        "da −0.069 a −0.154**, un peggioramento di 8 punti di IR. "
        "Statisticamente è dentro la banda CI, ma "
        "il senso è netto: aggiungere la band sopra un fit già "
        "liquidity-aware NON aiuta, anzi degrada lievemente "
        "il punto-stima dell'IR.",
    )
    add_para(
        doc,
        "Interpretazione: il fit Liquidity-aware è già parsimonioso "
        "(turnover 0.0045 è il più basso del notebook); la band non ha "
        "rebalance superflui da saltare. Quando occasionalmente \"hold\"-a "
        "i pesi vecchi invece di aggiornarli, blocca aggiornamenti "
        "informativi che il fit avrebbe fatto in modo già "
        "cost-conscious. Risultato: la replica perde aderenza al "
        "drift del target e l'IR cala.",
    )

    box = open_methodology_box(
        doc,
        "Perché Idea A non si può facilmente applicare a Kalman e NNLS "
        "(e quindi A+B esiste solo su base EN)"
    )
    box_para(
        box,
        "Idea A non è un wrapper che gira sopra un modello: è una "
        "scelta di fit. La funzione fit_liquidity_aware risolve un "
        "problema convesso del tipo min ‖y − Xw‖² + Σⱼ κⱼ |Δwⱼ| + Σⱼ "
        "λⱼ |Δwⱼ|^1.5. Per portarla:",
    )
    box_bullets(
        box,
        [
            ("Su NNLS: ",
             "servirebbe riscrivere il solver come QP non-negativa con "
             "penalty κ/λ. Fattibile (cvxpy o scipy SLSQP), ma è un "
             "nuovo modello, non un factory swap."),
            ("Su Fully-invested NNLS: ",
             "come NNLS più il vincolo Σw = 1. Stesso ordine di "
             "complessità."),
            ("Su Kalman: ",
             "molto più invasivo. Bisognerebbe far entrare κⱼ nella "
             "covarianza di processo Q (anisotropa per asset) o nel "
             "prior. Non è una factory swap, è un Kalman diverso."),
        ],
    )
    box_para(
        box,
        "Idea B invece è plug-and-play: il backtest_adaptive accetta un "
        "model_factory generico (per EN/NNLS) e il "
        "backtest_adaptive_kalman risolve il caso Kalman con la stateful "
        "che abbiamo discusso. È per questo che B vive su tutti i "
        "modelli mentre A vive solo su EN.",
    )

    # 7.4 Verdetto e confronto finale
    add_heading(doc, "7.4 Quadro finale dei production-oriented", level=2)

    fig = ROOT / "_report_figs" / "sec7_extended_cumulative_2.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Returns cumulati net-of-cost — tutte le varianti "
                    "della Sezione 7. Le sette curve sono "
                    "indistinguibili visivamente nel tratto principale: "
                    "la qualità di tracking è equivalente. Le "
                    "differenze sono nei dettagli di gross exposure, "
                    "turnover e IR puntuale.",
            width_cm=15,
        )
    fig = ROOT / "_report_figs" / "sec7_extended_cumulative_3.png"
    if fig.exists():
        add_figure(
            doc, fig,
            caption="Turnover settimanale medio: classifica orizzontale "
                    "delle varianti. Adaptive band (Kalman) è la più "
                    "parsimoniosa (0.0041), seguita da Liq + Adaptive "
                    "(EN) (0.0046) e Liquidity-aware da sola (0.0044). "
                    "Le tre vivono nello stesso ordine di grandezza, "
                    "molto sotto a EN tuned (0.0182).",
            width_cm=14,
        )

    add_para(
        doc,
        "Tabella di chiusura: TE/IR/Corr/Gross/Turnover per tutti i "
        "candidati production-oriented + i baseline più rilevanti delle "
        "sezioni precedenti.",
    )
    add_table(
        doc,
        header=["Modello", "TE net", "IR net", "Corr", "Gross exp.",
                "Turnover/sett."],
        rows=[
            ["EN tuned (riferimento)",        "0.0281", "−0.199", "0.853",
             "0.64", "0.0182"],
            ["NNLS long-only",                 "0.0275", "−0.169", "0.860",
             "0.64", "0.0160"],
            ["Fully-invested NNLS",            "0.0275", "−0.023", "0.860",
             "1.00", "0.0162"],
            ["Kalman (q = 10⁻⁷)",              "0.0276", "−0.154", "0.859",
             "0.48", "0.0077"],
            ["Idea 4 (Constrained)",           "0.0278", "−0.152", "0.856",
             "0.66", "0.0119"],
            ["**Liquidity-aware (Idea A)**",   "0.0276", "**−0.069**", "0.859",
             "0.77", "**0.0045**"],
            ["Adaptive band on EN (Idea B)",   "0.0280", "−0.203", "0.854",
             "0.63", "0.0137"],
            ["Adaptive band on Kalman",        "0.0279", "−0.146", "0.855",
             "0.48", "0.0041"],
            ["Adaptive band on NNLS",          "0.0275", "−0.161", "0.860",
             "0.63", "0.0121"],
            ["Liq + Adaptive (A+B)",           "0.0278", "−0.154", "0.857",
             "0.77", "0.0046"],
        ],
    )
    add_para(
        doc,
        "Il vincitore di produzione, alla luce di tutte le metriche "
        "viste, è **Liquidity-aware da solo**. Quattro ragioni:",
    )
    add_bullets(
        doc,
        [
            ("Miglior IR netto puntuale fra tutti i replicatori (escluso "
             "Fully-invested NNLS, che è un caso a parte — vedi sotto). ",
             "IR −0.069 è quasi metà di Kalman/Idea 4/NNLS+B."),
            ("Gross exposure 0.77, ",
             "il più alto fra i sani — più vicino al fully-invested "
             "che gli altri modelli a turnover basso. Riduce in modo "
             "tangibile il Beta gap (Sezione 4)."),
            ("Turnover 0.0045, ",
             "indistinguibile dai migliori (Kalman+B 0.0041, A+B "
             "0.0046). Niente da guadagnare passando alle varianti "
             "con la band."),
            ("Correlation 0.859, ",
             "in linea con i migliori modelli senza band."),
        ],
    )
    add_para(
        doc,
        "Fully-invested NNLS resta un'opzione alternativa: il suo IR "
        "(−0.023) è ancora più vicino a zero del Liquidity-aware, ma "
        "lo paga in gross exposure = 1.00 (al limite UCITS, niente "
        "spazio per coperture) e turnover 0.016 (3.5× peggio). Liquidity-"
        "aware vince sul compromesso operativo complessivo; "
        "Fully-invested NNLS vince se l'unica priorità è azzerare il "
        "bias.",
    )

    # 7.5 Integrazione con idee precedenti
    add_heading(doc, "7.5 Integrazione con i modelli delle sezioni "
                     "precedenti", level=2)
    add_para(
        doc,
        "Domanda esplicita: A e B si possono comporre anche con i "
        "modelli \"buoni\" delle sezioni precedenti? Rispondiamo "
        "distinguendo cosa è plug-and-play da cosa richiede una "
        "riscrittura.",
    )

    add_heading(doc, "Quello che già esiste e si può combinare a costo zero",
                level=3)
    add_bullets(
        doc,
        [
            ("Adaptive band (B) ⊕ qualunque modello con factory "
             "sklearn-style: ",
             "EN, Lasso, Ridge, NNLS long-only, NNLS Fully-invested, "
             "Idea 4 Constrained, Idea 1 Multi-moment. Si passa "
             "semplicemente il factory a backtest_adaptive(). "
             "Già fatto per EN/Kalman/NNLS — manca all'appello "
             "Fully-invested NNLS e Idea 4 Constrained."),
            ("Adaptive band ⊕ Fully-invested NNLS — scelta interessante: ",
             "ci si aspetta turnover ~50% di Fully-invested NNLS "
             "(0.016 → ~0.010), e IR sempre vicino a zero. È la "
             "combinazione che unisce \"closure del Beta gap di "
             "livello\" + \"meno rebalance\" — quindi probabilmente la "
             "miglior versione del Fully-invested NNLS in produzione."),
            ("Adaptive band ⊕ Idea 4 (Constrained): ",
             "Idea 4 ha già un turnover modesto (0.012), la band "
             "potrebbe scenderlo a 0.008-0.009 senza perdere "
             "interpretabilità dei vincoli macro. Combinazione naturale "
             "per un prodotto regolamentato."),
        ],
    )

    add_heading(doc, "Quello che richiede una riscrittura del modello",
                level=3)
    add_bullets(
        doc,
        [
            ("Liquidity-aware (A) ⊕ NNLS long-only: ",
             "fattibile in 1-2 ore di lavoro. Si riscrive il solver "
             "come QP convessa con vincoli wⱼ ≥ 0 e penalty per-asset "
             "κⱼ|Δwⱼ| + λⱼ|Δwⱼ|^1.5. Risultato atteso: turnover "
             "ulteriormente ridotto rispetto al solo A (0.004 → "
             "~0.003), e gross verosimilmente meno alto (perché "
             "non-negatività + penalty κ congiuntamente disincentivano "
             "le posizioni)."),
            ("Liquidity-aware (A) ⊕ Fully-invested NNLS (Σw = 1): ",
             "come sopra più il vincolo Σw=1. Il risultato sarebbe "
             "potenzialmente il \"campione assoluto\": Beta = 1 by "
             "design + turnover minimizzato + costi per-asset "
             "minimizzati. Da provare seriamente."),
            ("Liquidity-aware (A) ⊕ Kalman: ",
             "molto più invasivo. Bisognerebbe far entrare κⱼ nella "
             "covarianza di processo Q (rendendola anisotropa: futures "
             "illiquidi → Q smaller → pesi più stabili) o nel prior "
             "iniziale. È un Kalman diverso, non un wrapper."),
            ("Idea 5 (turnover penalty L1 in spazio reale) ⊕ "
             "Liquidity-aware (A): ",
             "le due penalty sono parenti — entrambe colpiscono il "
             "turnover, ma con metriche diverse (λ_TO costo uniforme "
             "vs κⱼ costo per-asset). Combinarle potrebbe sembrare "
             "ridondante, ma offrirebbe un controllo separato sul "
             "turnover (Idea 5) e sul mix di asset (Idea A). Vale la "
             "pena testarlo se prima Idea 5 viene integrata nel "
             "backtest engine unificato (vedi note di Sezione 6)."),
        ],
    )

    add_heading(doc, "Quello che già è stato considerato ma con risultati "
                     "deboli", level=3)
    add_bullets(
        doc,
        [
            ("Multi-moment matching (Idea 1) ⊕ Liquidity-aware: ",
             "Idea 1 non funziona bene (Sezione 5) — il moment "
             "matching peggiora con λ crescente, l'optimizer overshoot. "
             "Aggiungere la liquidity penalty non risolve quel "
             "problema strutturale. Da scartare."),
            ("HMM regime-aware ⊕ Adaptive band: ",
             "interessante in linea di principio: regime detector + "
             "skip rebalances in regimi stabili. Ma HMM oggi non è "
             "validato (Sezione 5), e quindi la combinazione resta una "
             "direzione di lavoro futuro."),
        ],
    )

    add_heading(doc, "Sintesi: le tre combinazioni che vale la pena "
                     "testare nella prossima iterazione", level=3)
    add_bullets(
        doc,
        [
            ("Priorità 1: Adaptive band ⊕ Fully-invested NNLS. ",
             "Costo zero (factory già esistente). Aspettativa: turnover "
             "0.016 → 0.010, IR resta vicino a zero, gross resta 1. È "
             "il candidato più forte per la consegna se vince il "
             "Liquidity-aware su uno o due piccoli dettagli."),
            ("Priorità 2: Liquidity-aware ⊕ Fully-invested NNLS. ",
             "Costo ~2 ore. Aspettativa: il \"campione assoluto\" — "
             "Beta=1, turnover minimo, mix di asset cost-aware. "
             "Se funziona, è il candidato finale di produzione."),
            ("Priorità 3: Adaptive band ⊕ Idea 4 (Constrained). ",
             "Costo zero. Aspettativa: stessa interpretabilità di Idea "
             "4 ma con un terzo di turnover in meno. Soluzione "
             "produttiva con la storia più semplice da raccontare a un "
             "PM."),
        ],
    )

    # 7.6 Flags
    flags = [
        ("Liquidity-aware (Idea A) è il vincitore di produzione del "
         "notebook attuale. ",
         "IR puntuale −0.07 (migliore di tutti), turnover 0.0045 (più "
         "basso fra i sani), gross 0.77 (più alto fra i sani). Da "
         "consigliare per la consegna salvo che si decida di "
         "implementare le combinazioni della Priorità 2 (A ⊕ "
         "Fully-invested NNLS)."),
        ("Adaptive band (Idea B) NON è un vincitore strutturale. ",
         "Riduce il turnover di ~25-50% sui modelli a cui viene "
         "applicata, ma TE/IR/correlation cambiano dentro la banda CI "
         "bootstrap. A 2 bps di costo, il risparmio è 2-9 bps/anno: "
         "tangibile, non decisivo."),
        ("A+B (combinato) è statisticamente equivalente al solo A, ",
         "ma punto-stima leggermente peggiore. Non c'è ragione di "
         "preferirlo al solo A salvo che ci si aspetti un futuro shift "
         "di regime in cui la band dia valore aggiunto."),
        ("Idea A è fit-level e si compone solo con modelli "
         "regression-style. ",
         "Per portarla su NNLS/Kalman serve riscrivere il solver. Idea "
         "B è policy-level e si compone con tutti."),
        ("Tre combinazioni promettenti non ancora testate: ",
         "(i) B ⊕ Fully-invested NNLS, (ii) A ⊕ Fully-invested NNLS, "
         "(iii) B ⊕ Idea 4 Constrained. La (ii) è probabilmente la "
         "più forte ma ha costo implementativo."),
        ("Il muro R² ≤ 78.6% rimane il limite per tutte queste "
         "estensioni. ",
         "Liquidity-aware vince nella Beta CI bootstrap il quanto "
         "possibile sotto quel limite (Beta media ~0.78-0.80, "
         "indistinguibile dal teorico max 0.886)."),
        ("Bug latente di TC_BPS in cella 66 ora reso innocuo da "
         "defensive resets in celle 73/75/77. ",
         "Per le future iterazioni del notebook, considerare di "
         "rinominare la variabile in Idea 5 (es. TC_BPS_IDEA5) per "
         "rimuovere completamente la fonte del problema invece di "
         "patcharne i sintomi a valle."),
    ]
    add_flag_box(doc, "Flag per il prosieguo del progetto", flags)

    add_para(doc, "")
    return doc


# ---------- SECTION 8: Forward-looking diagnostics ---------------------------

def section_8(doc):
    add_heading(doc, "8. Diagnostiche forward-looking: stress test, "
                     "conformal prediction, RL rebalancing, replay "
                     "storico", level=1)
    add_para(
        doc,
        "Tutte le sezioni precedenti hanno costruito e validato modelli "
        "di replica. Questa sezione descrive quattro strumenti che, "
        "applicati a un modello già scelto, rispondono a domande di "
        "rischio e operatività non più sulla qualità del fit ma sulla "
        "sua usabilità in produzione. Il blocco è stato disegnato "
        "deliberatamente come model-agnostic: c'è un selettore "
        "centralizzato che identifica il replicatore corrente, e tutte "
        "le idee a valle lo consumano via lo stesso nome di variabile. "
        "Cambiando la riga del selettore si rieseguono le quattro idee "
        "su un modello diverso senza toccare altro.",
    )
    add_para(
        doc,
        "Importante per il report: questa sezione descrive cosa "
        "ciascuna idea fa, con quale metodo, e come si integra con il "
        "resto del notebook. I numeri concreti dei risultati non "
        "vengono riportati qui: ogni idea va eseguita sul modello di "
        "produzione finale (in corso di decisione fra Liquidity-aware "
        "e le combinazioni della Sezione 7.5) e i suoi risultati "
        "andranno presentati come capitolo a sé nella consegna finale.",
    )

    # 8.0 chosen_model
    add_heading(doc, "8.1 Selettore del modello — chosen_model", level=2)
    add_para(
        doc,
        "All'inizio del blocco forward-looking c'è un'unica cella di "
        "configurazione che definisce quattro variabili globali che "
        "verranno consumate da tutte le idee successive:",
    )
    add_bullets(
        doc,
        [
            ("chosen_model: ",
             "il dict prodotto da backtest() che contiene returns, "
             "pesi, gross exposure e VaR del modello scelto. Usato "
             "dalle Idee 8 (per estrarre i pesi finali da stressare), "
             "8b (per il replay storico), e indirettamente da D ed E."),
            ("chosen_factory: ",
             "la factory sklearn-style del modello (es. best_factory = "
             "lambda: ElasticNet(alpha=best_alpha, l1_ratio=best_l1_ratio)). "
             "Usata da Idea D (conformal prediction) e Idea E (RL "
             "environment) per refittare il modello nel rolling loop. "
             "Caveat: Kalman e modelli ad hoc come Liquidity-aware non "
             "hanno una factory sklearn-style e per loro bisogna "
             "passare best_factory (EN) o gestire un wrapper "
             "dedicato."),
            ("chosen_weights: ",
             "ultima riga di chosen_model['weights'] — la snapshot "
             "corrente del portafoglio, usata da Idea 8 (per il VaR "
             "condizionato) e 8b (per il replay storico)."),
            ("chosen_net_ret: ",
             "i returns OOS netti del modello scelto, usati come "
             "baseline di confronto in tutte le idee successive."),
        ],
    )
    add_para(
        doc,
        "Default attuale: chosen_model = elnet_monthly (Elastic Net "
        "tuned). Dato il verdetto della Sezione 7 (Liquidity-aware "
        "vincitore), il candidato naturale per la consegna è "
        "chosen_model = liq_result. La Sezione 9 traccia esplicitamente "
        "questo come Priorità 1 del piano di consolidamento.",
    )

    # 8.2 Idea 8
    add_heading(doc, "8.2 Idea 8 — Inverse stress testing via copula "
                     "t-Student", level=2)
    add_heading(doc, "Cosa fa", level=3)
    add_para(
        doc,
        "Risponde alla domanda \"se i futures si muovessero "
        "congiuntamente in uno scenario plausibile ma estremo, quanto "
        "perderebbe la replica?\". Lo fa generando 10 000 scenari "
        "settimanali sintetici che preservano sia le distribuzioni "
        "marginali dei singoli futures sia la struttura di dipendenza "
        "fra loro, e poi calcolando la distribuzione del P&L della "
        "replica condizionato a combinazioni specifiche di rendimenti "
        "di equity e bond.",
    )
    add_heading(doc, "Come funziona", level=3)
    add_bullets(
        doc,
        [
            ("Fit della copula t-Student: ",
             "Per ogni coppia di futures si stima il tau di Kendall "
             "e lo si converte in correlazione di Spearman via "
             "ρ = sin(π·τ/2). La matrice di correlazione "
             "risultante viene resa positiva-definita aggiungendo un "
             "piccolo jitter agli autovalori. Il grado di libertà ν "
             "viene scelto con profile maximum likelihood su una "
             "griglia ν ∈ {3, ..., 30}."),
            ("Generazione degli scenari: ",
             "Si campionano vettori multinormali z ~ N(0, ρ), si "
             "scalano per √(χ²_ν/ν) ottenendo realizzazioni "
             "multivariate t-Student, e si trasformano in "
             "pseudo-osservazioni U ∈ [0, 1]^11 via la CDF univariata "
             "t. Da U a returns: per ogni asset si applica la CDF "
             "empirica inversa dei suoi returns storici. Risultato: "
             "10 000 scenari (un vettore di 11 returns ciascuno) che "
             "rispettano sia le code marginali sia la struttura di "
             "dipendenza tail."),
            ("VaR condizionato: ",
             "Sui 10 000 scenari si calcola il P&L della replica come "
             "scenari @ chosen_weights. Si condiziona su griglie di "
             "equity return × bond return e per ogni cella si estrae "
             "il 5° percentile del P&L. Risultato: una heat map "
             "VaR(equity_return, bond_return)."),
        ],
    )
    add_heading(doc, "Cosa produce", level=3)
    add_bullets(
        doc,
        [
            ("Heat map del VaR condizionato",
             ": rossa nei quadranti equity-down/bond-down (worst-case)."),
            ("KDE della distribuzione del P&L della replica",
             "con marcatori dei percentili 5°/50°/95°."),
            ("Box plot del P&L condizionato",
             "a diverse combinazioni equity/bond."),
        ],
    )

    # 8.3 Idea D
    add_heading(doc, "8.3 Idea D — Conformal prediction sulla TE futura "
                     "(con D2 regime-conditioned)", level=2)
    add_heading(doc, "Cosa fa", level=3)
    add_para(
        doc,
        "Risponde alla domanda \"con quale probabilità il prossimo "
        "tracking deviation cadrà entro X bps?\" — una domanda che "
        "intervalli di confidenza bootstrap (Sezione 4) NON "
        "rispondono, perché quelli sono retrospettivi e assumono "
        "stazionarietà. La conformal prediction è "
        "distribution-free e fornisce una garanzia di copertura "
        "marginale: dato un livello α (es. 0.10), l'intervallo "
        "prodotto contiene il valore realizzato con probabilità "
        "almeno 1−α, senza assumere distribuzione del rumore.",
    )
    add_heading(doc, "Come funziona (split conformal walk-forward)",
                level=3)
    add_bullets(
        doc,
        [
            ("Schema temporale per ogni step di rebalance: ",
             "si divide la finestra rolling in due blocchi cronologici "
             "— train_set (settimane più lontane) e cal_set "
             "(settimane più recenti, 52 settimane = 1 anno per "
             "default). Il modello viene fittato sul train_set, e si "
             "calcolano i residui |y − ŷ| sul cal_set."),
            ("Quantile dei residui: ",
             "q = quantile (1−α)(n+1)/n dei nonconformity scores. "
             "La correzione (n+1)/n è la versione finite-sample "
             "garantita."),
            ("Intervallo per la settimana successiva: ",
             "[X_t @ β + α_intercetta − q, X_t @ β + α_intercetta + q]. "
             "Validazione: si misura empiricamente la coverage rate "
             "sull'OOS — dovrebbe convergere a (1−α)."),
        ],
    )
    add_heading(doc, "Variante D2 — regime-conditioned", level=3)
    add_para(
        doc,
        "Il quantile q standard è invariante nel tempo. D2 lo "
        "stratifica per regime: si classificano le settimane di "
        "calibrazione in calm vs stress (sopra/sotto la mediana della "
        "vol rolling 26 settimane del target), si calcola un quantile "
        "separato q_calm e q_stress, e si applica quello "
        "corrispondente al regime corrente. Risultato: intervalli più "
        "stretti in calm e più larghi in stress — informazione utile "
        "per il risk management. La garanzia di copertura passa da "
        "marginale (uniforme su tutto il sample) a condizionale "
        "(condizionata al regime), che è più informativa quando il "
        "regime è osservabile.",
    )
    add_heading(doc, "Cosa produce", level=3)
    add_bullets(
        doc,
        [
            ("Coverage rate empirica vs nominale: ",
             "diagnostica della validità del metodo."),
            ("Banda dei prediction intervals nel tempo: ",
             "si vede come l'incertezza si allarga in stress e si "
             "stringe in calm."),
            ("Confronto D vs D2: ",
             "metà-larghezza media del CI in calm vs stress; "
             "se D2 fornisce bande significativamente diverse nei "
             "due regimi, vince in informativeness."),
        ],
    )

    # 8.4 Idea E
    add_heading(doc, "8.4 Idea E — Reinforcement learning per la policy di "
                     "rebalancing (con E2 copula-augmented)", level=2)
    add_heading(doc, "Cosa fa", level=3)
    add_para(
        doc,
        "Rilegge Idea B (adaptive band) come problema di decision "
        "policy e lo affronta col Q-learning tabulare. La domanda "
        "diventa: \"dato lo stato corrente di mercato — vol, drift dei "
        "pesi rispetto all'ideale, TE recente — l'azione ottimale è "
        "hold, partial rebalance o full rebalance?\". Idea B "
        "rispondeva con una regola euristica; Idea E lascia che il "
        "trade-off venga imparato dai dati.",
    )
    add_heading(doc, "Setup RL", level=3)
    add_bullets(
        doc,
        [
            ("Stato: ",
             "(vol_regime, weight_drift, recent_TE) discretizzati su 3 "
             "livelli ciascuno → 27 stati totali. vol_regime usa σ del "
             "target su 26 settimane, drift = Σ|w − ideal_w|, recent_TE "
             "su finestra recente."),
            ("Azioni: ",
             "0 = hold (tieni i pesi correnti), 1 = partial rebalance "
             "(50% verso ideal), 2 = full rebalance."),
            ("Reward: ",
             "−(TE² + λ · costo), dove λ è un peso che bilancia "
             "tracking quality vs costi di transazione (default = "
             "1.0)."),
            ("Algoritmo: ",
             "Q-learning tabulare con ε-greedy decaying (ε passa da "
             "0.3 a 0.05), learning rate 0.1, discount γ = 0.95, 300 "
             "episodi ciascuno full-sample walk-through."),
            ("Pre-compute: ",
             "i pesi ideali w_ideal[i] e il loro intercetta a_ideal[i] "
             "sono pre-calcolati su tutto il sample una volta sola "
             "(fit del chosen_factory + VaR scaling) e cached, per "
             "rendere ciascun episodio veloce. Questa è la ragione "
             "per cui Idea E ha bisogno di un chosen_factory "
             "sklearn-compatibile, non solo di un dict di pesi."),
        ],
    )
    add_heading(doc, "Variante E2 — copula-augmented training", level=3)
    add_para(
        doc,
        "Il limite dell'RL standard è che impara su una sola "
        "traiettoria storica (~500 settimane). Idea E2 affianca a "
        "questa training data scenari sintetici generati dalla copula "
        "di Idea 8: rendono il training set più ricco di code, "
        "permettendo all'agente di vedere combinazioni di stress che "
        "il sample reale non contiene in numero sufficiente. "
        "L'agente viene poi valutato sulla traiettoria reale OOS.",
    )
    add_heading(doc, "Cosa produce", level=3)
    add_bullets(
        doc,
        [
            ("Convergenza della learning curve",
             "(reward medio per episodio nel tempo)."),
            ("Mappa di policy ottimale",
             "sui 27 stati: cosa fa l'agente in ogni combinazione "
             "vol/drift/TE."),
            ("Confronto con policy fisse",
             "(rebalance ogni 1/4/12 settimane): RMSE della TE, "
             "costi di transazione totali, % settimane in cui "
             "l'agente rebalancia."),
            ("Confronto E vs E2",
             "sulla stessa metrica, per quantificare il valore "
             "aggiunto degli scenari sintetici."),
        ],
    )

    # 8.5 Idea 8b
    add_heading(doc, "8.5 Idea 8b — Replay di scenari storici", level=2)
    add_heading(doc, "Cosa fa", level=3)
    add_para(
        doc,
        "Complementare a Idea 8: invece di generare scenari "
        "sintetici, replaya sei crisi storiche reali (2008 GFC, 2010 "
        "Flash Crash, 2011 EU Debt, 2015 China Sell-off, 2018 Q4, "
        "2020 COVID) attraverso il chosen_model e confronta il "
        "P&L della replica con quello del target, settimana per "
        "settimana. Domanda: in quanto la replica avrebbe "
        "amplificato o ammortizzato la perdita?",
    )
    add_heading(doc, "Come funziona — con la cautela del look-ahead",
                level=3)
    add_para(
        doc,
        "Implementazione: si applicano i pesi più recenti del "
        "chosen_model (chosen_weights = "
        "chosen_model['weights'].iloc[−1]) ai returns storici dei "
        "futures durante ciascuna crisi. Per ogni crisi si "
        "calcolano: P&L cumulato del target, P&L cumulato della "
        "replica, divergenza massima settimanale, beta della "
        "replica vs target sulla finestra, correlazione, max "
        "drawdown.",
    )
    add_para(
        doc,
        "Avvertenza esplicita di look-ahead bias: i pesi usati sono "
        "quelli al ~2021, applicati retroattivamente a una crisi del "
        "2008. Riflettono la struttura di mercato corrente, non "
        "quella all'epoca della crisi. L'interpretazione corretta è "
        "\"come si sarebbe comportata la replica corrente in uno "
        "scenario simile a quella crisi\", NON \"come si sarebbe "
        "comportata una replica live al tempo della crisi\". Per la "
        "seconda lettura servirebbero pesi rolling stimati in-sample "
        "a ogni crisi — direzione di lavoro futuro.",
    )
    add_heading(doc, "Cosa produce", level=3)
    add_bullets(
        doc,
        [
            ("Tabella sintetica per crisi",
             ": target cum, replica cum, divergenza max, β, "
             "correlazione."),
            ("Bar chart del max drawdown",
             "per ciascuna crisi, target vs replica."),
            ("Beta per-crisi",
             "(slope ols replica su target nella finestra): "
             "diagnostica di come la replica si scala in ciascun "
             "regime di stress storico."),
        ],
    )

    # Closing
    add_para(
        doc,
        "Nessun risultato numerico è incluso in questa sezione: tutte "
        "e quattro le idee sono diagnostiche del modello di produzione "
        "finale, che andrà deciso al termine del consolidamento "
        "descritto in Sezione 9. Quando il chosen_model sarà fissato "
        "(probabilmente liq_result, o una delle combinazioni di "
        "Sezione 7.5), questa sezione andrà rieseguita e i suoi "
        "output presentati in modo organico — sia come validazione "
        "del rischio di portafoglio (Idee 8 / 8b / D / D2), sia come "
        "componente operativa potenziale (Idea E / E2 al posto di "
        "una policy fissa di rebalance).",
    )
    return doc


# ---------- SECTION 9: Roadmap finale ----------------------------------------

def section_9(doc):
    add_heading(doc, "9. Roadmap finale per la coerenza del progetto",
                level=1)
    add_para(
        doc,
        "Le otto sezioni precedenti hanno costruito tutti i mattoni "
        "del progetto in modo tecnicamente solido, ma il progetto nel "
        "suo insieme contiene ancora una serie di discontinuità — "
        "punti in cui un risultato di una sezione non è stato "
        "ancora portato a valle, o in cui un modello secondario "
        "viene usato dove si potrebbe usare il modello vincitore. "
        "Questa sezione finale raccoglie tutte queste discontinuità "
        "in un piano d'azione ordinato per importanza, attingendo "
        "esplicitamente dalle flag che si sono accumulate sezione "
        "per sezione.",
    )

    # P1
    add_heading(doc, "Priorità 1 — Spostare il \"production model\" "
                     "dalle Idee 8 / 8b / D / E", level=2)
    add_para(
        doc,
        "Tutte e quattro le idee forward-looking della Sezione 8 oggi "
        "consumano elnet_monthly come default (via chosen_model). "
        "La Sezione 4 ha mostrato che EN tuned è statisticamente "
        "dominato sul Beta; la Sezione 7 ha eletto Liquidity-aware "
        "come vincitore di produzione. Le Idee 8 / 8b / D / E "
        "dovrebbero rifletterlo.",
    )
    add_bullets(
        doc,
        [
            ("Cambio operativo: ",
             "in cella 81 sostituire chosen_model = elnet_monthly con "
             "chosen_model = liq_result (oppure con il vincitore "
             "delle combinazioni di Priorità 2). chosen_weights e "
             "chosen_net_ret seguono automaticamente."),
            ("Caveat factory: ",
             "Liquidity-aware non ha una factory sklearn-style. "
             "Per Idea D e Idea E che richiedono il refit nel "
             "rolling loop, due opzioni: (a) creare una "
             "LiquidityAwareWrapper sklearn-compatible che applichi "
             "fit_liquidity_aware in .fit() — ~1 ora di lavoro; (b) "
             "mantenere best_factory (EN) come chosen_factory per "
             "D/E, accettando l'asimmetria. La (a) è più pulita."),
            ("Conseguenza per il report: ",
             "Idea 8 e 8b vanno presentate come stress test del "
             "modello reale; le bande di Idea D diventano la quantile "
             "forward del Liquidity-aware, statisticamente più "
             "informativa di quella di EN tuned."),
        ],
    )

    # P2
    add_heading(doc, "Priorità 2 — Testare le tre combinazioni della "
                     "Sezione 7.5", level=2)
    add_para(
        doc,
        "Le combinazioni identificate ma non ancora valutate "
        "numericamente:",
    )
    add_bullets(
        doc,
        [
            ("Priorità 2A: Adaptive band ⊕ Fully-invested NNLS. ",
             "Costo zero (factory già esistente: "
             "fully_invested_nnls_factory). Aspettativa: turnover "
             "0.016 → ~0.010, IR resta vicino a zero (chiusura del "
             "Beta gap di livello), gross resta 1. Se vince, sostituisce "
             "Liquidity-aware come vincitore."),
            ("Priorità 2B: Liquidity-aware ⊕ Fully-invested NNLS "
             "(Σw = 1). ",
             "Costo ~2 ore: estendere fit_liquidity_aware con vincolo "
             "Σw = 1 (SLSQP o cvxpy). Aspettativa: il \"campione "
             "assoluto\" — Beta = 1 by design + turnover minimo + "
             "costi per-asset minimizzati."),
            ("Priorità 2C: Adaptive band ⊕ Idea 4 (Constrained). ",
             "Costo zero (factory già esistente: constrained_factory). "
             "Aspettativa: stessa interpretabilità di Idea 4 ma con un "
             "terzo di turnover in meno. Il candidato con la \"storia\" "
             "più semplice da raccontare a un PM."),
        ],
    )

    # P3
    add_heading(doc, "Priorità 3 — Integrare Idea 5 nel backtest engine "
                     "unificato", level=2)
    add_para(
        doc,
        "Idea 5 (Elastic Net + turnover penalty) vive oggi in una "
        "cella isolata con parametri propri (window 78, TC 5 bps, "
        "loop custom). Questo rende i suoi numeri non direttamente "
        "comparabili con il resto del notebook. La Sezione 6 ha "
        "già documentato tre livelli di intervento; in ordine di "
        "costo:",
    )
    add_bullets(
        doc,
        [
            ("Livello 1 (mezz'ora): ",
             "riallineare window=208, TC_BPS=2e-4, aggiungere VaR "
             "Cornish-Fisher con cap UCITS. Numeri di Idea 5 "
             "comparabili al resto."),
            ("Livello 2 (significativo): ",
             "wrap completo come model factory sklearn-style con "
             "stato interno (analogo a backtest_adaptive_kalman). La "
             "turnover penalty diventa un ingrediente trasferibile a "
             "Ridge, Lasso, NNLS, Kalman — esattamente come la "
             "Liquidity-aware penalty di Idea A."),
            ("Livello 3 (ambizioso): ",
             "libreria di penalty componibili (L1, L2, turnover, "
             "liquidity). Permette combinazioni come Idea 5 + Idea A "
             "che oggi non sono testabili."),
        ],
    )

    # P4
    add_heading(doc, "Priorità 4 — Validare HMM (Idea 2)", level=2)
    add_para(
        doc,
        "HMM è descritto e implementato (Sezione 5.2), il design è "
        "solido (feature causali, no leakage, regimi indipendenti dal "
        "target), ma l'esecuzione attuale non produce risultati "
        "stampabili. Quando sarà ri-eseguibile end-to-end, validare "
        "tre cose:",
    )
    add_bullets(
        doc,
        [
            ("Posterior P(stress) deve allinearsi visivamente ai "
             "drawdown noti: ",
             "2008-09, 2011, 2015-16, 2018 Q4, marzo 2020."),
            ("Pesi per-regime devono avere senso economico: ",
             "calm → più equity, stress → più bond/oro (rotazione "
             "flight-to-quality)."),
            ("Confronto 2 vs 3 regimi: ",
             "n/p ≈ 9 con 2 regimi, n/p ≈ 6 con 3. Sui nostri dati 2 "
             "regimi dovrebbero battere 3."),
        ],
    )

    # P5
    add_heading(doc, "Priorità 5 — Estendere il bootstrap CI a tutti i "
                     "nuovi modelli", level=2)
    add_para(
        doc,
        "Sezione 4 ha calcolato CI bootstrap su 8 modelli (i baselines "
        "+ EN tuned, NNLS, Fully-invested NNLS, Kalman). Le sezioni "
        "successive (6 e 7) hanno aggiunto Idea 4 Constrained, Idea 5, "
        "Liquidity-aware, Adaptive band su 3 modelli, Liq + Adaptive. "
        "Tutti questi modelli mancano dal bootstrap. Estendere il "
        "bootstrap a loro permetterebbe di dire con confidenza quando "
        "due modelli sono statisticamente diversi e quando sono "
        "indistinguibili. Costo: una loop chiamando stationary_bootstrap "
        "(già implementata) su ogni nuovo modello.",
    )

    # P6
    add_heading(doc, "Priorità 6 — Rinominare TC_BPS in Idea 5", level=2)
    add_para(
        doc,
        "La fix definitiva del bug latente che produceva i grafici "
        "rotti delle adaptive band. Oggi è patchato in difensiva "
        "(reset di TC_BPS = 2e-4 in testa alle celle 73, 75, 77), "
        "ma la causa root resta in cella 66 (TC_BPS = 5 inteso come "
        "bps integer). Rinominare a TC_BPS_IDEA5 = 5 nella cella di "
        "Idea 5 elimina la pollution del namespace globale. È un "
        "intervento di una riga, ma chiude il debito tecnico.",
    )

    # P7
    add_heading(doc, "Priorità 7 — Forensic alpha decomposition", level=2)
    add_para(
        doc,
        "L'IR floor a ~−0.15 osservato in tutti i modelli (Sezione 3 "
        "flag) rappresenta circa 50 bps/anno di drift non replicabile. "
        "L'Idea 6 nel recap originale del notebook proponeva di "
        "decomporre questo residuo: quanto è dovuto al smoothing "
        "strutturale di HFRX, quanto a posizioni illiquide degli "
        "hedge fund, quanto a manager skill. È un esercizio "
        "diagnostico, non un nuovo modello, ma chiude il cerchio "
        "interpretativo: \"il nostro replicatore lascia 50 bps di "
        "alpha sul tavolo — ecco la sua composizione\".",
    )

    # P8
    add_heading(doc, "Priorità 8 — Estendere l'universo di replica",
                level=2)
    add_para(
        doc,
        "Il muro R² ≤ 78.6% sul Monster Index (Sezione 4) non si "
        "supera con miglioramenti del modello: si supera aggiungendo "
        "regressori. Tre direzioni naturali:",
    )
    add_bullets(
        doc,
        [
            ("CDS futures e/o index ETFs di credito: ",
             "catturano lo spread component che LEGATRUU (R² = 0.51) "
             "ha e i 4 bond futures attuali no."),
            ("VIX futures: ",
             "catturano la vol-of-vol che HFRX implicitamente "
             "trada via posizioni opzionali (R² HFRX = 0.27)."),
            ("EM bond futures e currency carry: ",
             "il resto del span che manca a LEGATRUU."),
        ],
    )
    add_para(
        doc,
        "È probabilmente fuori scope per il presente progetto, ma è "
        "la naturale prossima iterazione e va menzionata come limite "
        "esplicito.",
    )

    # P9
    add_heading(doc, "Priorità 9 — Considerare la stabilità della "
                     "snapshot finale dei pesi", level=2)
    add_para(
        doc,
        "La Sezione 2.7 ha segnalato che i pesi finali di EN tuned "
        "hanno DU1 = +0.808 contro una media di −0.074: la singola "
        "settimana finale è atipica. Le Idee 8 e 8b usano "
        "chosen_weights = weights.iloc[−1], quindi sono "
        "potenzialmente sensibili a quella settimana. Due "
        "mitigazioni possibili:",
    )
    add_bullets(
        doc,
        [
            ("Smoothing dei pesi recenti: ",
             "media esponenziale degli ultimi N pesi (es. N = 13 "
             "settimane) come input alle idee diagnostiche. Riduce "
             "la sensibilità all'ultimo punto."),
            ("Sensitivity check: ",
             "rieseguire Idea 8 con tre snapshot di pesi diverse "
             "(es. ultime 4, 13, 26 settimane medie). Se i "
             "risultati cambiano poco, la sensibilità è OK; "
             "altrimenti il VaR riportato è artefatto."),
        ],
    )

    # P10
    add_heading(doc, "Priorità 10 — Tuning della sparsity threshold di "
                     "NNLS", level=2)
    add_para(
        doc,
        "Sparsity_thresh = 0.005 in NNLSWrapper e in "
        "FullyInvestedNNLS è fissato a mano (Sezione 3.1 flag). Non "
        "è ottimizzato. Una piccola griglia (es. 0.001, 0.005, "
        "0.01, 0.02, 0.05) costa poche righe e potrebbe migliorare "
        "marginalmente la stabilità delle composizioni. Bassa "
        "priorità, ma chiude un'altra flag aperta.",
    )

    # Closing
    add_para(doc, "")
    add_heading(doc, "Sintesi visiva delle priorità", level=2)
    add_table(
        doc,
        header=["Priorità", "Intervento", "Costo", "Valore"],
        rows=[
            ["1",  "chosen_model = liq_result + LiquidityAwareWrapper "
                  "per Idea D/E",
             "~1 h", "Alto"],
            ["2A", "B ⊕ Fully-invested NNLS",        "0",    "Alto"],
            ["2B", "A ⊕ Fully-invested NNLS",        "~2 h", "Alto (potenziale campione)"],
            ["2C", "B ⊕ Idea 4 Constrained",         "0",    "Medio"],
            ["3",  "Integrare Idea 5 nel backtest "
                  "engine unificato",
             "0.5–4 h", "Medio"],
            ["4",  "Validare HMM (Idea 2)",          "—",    "Medio"],
            ["5",  "Estendere bootstrap CI ai nuovi "
                  "modelli",
             "~30 min", "Medio"],
            ["6",  "Rinominare TC_BPS in cella 66",  "1 riga", "Tecnico"],
            ["7",  "Forensic alpha decomposition",   "~2 h", "Interpretativo"],
            ["8",  "Estendere l'universo (CDS, "
                  "VIX, EM)",
             "Fuori scope", "Alto (futuro)"],
            ["9",  "Smoothing della snapshot finale",
             "~30 min", "Medio"],
            ["10", "Tuning sparsity_thresh NNLS",
             "~30 min", "Basso"],
        ],
    )

    add_para(
        doc,
        "Eseguendo le Priorità 1 e 2 (le tre combinazioni della "
        "Sezione 7.5) si chiude il loop principale del progetto: il "
        "modello vincitore di produzione viene scelto in modo "
        "esplicito e tutte le diagnostiche forward-looking (Sezione "
        "8) si applicano a quel modello, non a un suo predecessore "
        "intermedio. Le altre priorità sono incrementali e possono "
        "essere affrontate in qualsiasi ordine.",
    )

    return doc


SECTIONS = {
    "1": section_1,
    "2": section_2,
    "3": section_3,
    "4": section_4,
    "5": section_5,
    "6": section_6,
    "7": section_7,
    "8": section_8,
    "9": section_9,
}


def main():
    if len(sys.argv) < 2:
        print("usage: py _report_builder.py <section_id|--rebuild-all>")
        print("available sections:", list(SECTIONS))
        sys.exit(1)
    arg = sys.argv[1]
    if arg == "--rebuild-all":
        if DOCX_PATH.exists():
            DOCX_PATH.unlink()
        doc = ensure_document()
        for sid, fn in SECTIONS.items():
            fn(doc)
            print(f"  + section {sid}")
        doc.save(DOCX_PATH)
        print(f"rebuilt all sections → {DOCX_PATH}")
        return
    sid = arg
    if sid not in SECTIONS:
        print(f"unknown section {sid!r}; available: {list(SECTIONS)}")
        sys.exit(1)
    doc = ensure_document()
    SECTIONS[sid](doc)
    doc.save(DOCX_PATH)
    print(f"wrote section {sid} → {DOCX_PATH}")


if __name__ == "__main__":
    main()
